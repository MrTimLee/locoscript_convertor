"""
Output converters: plain text, RTF, and DOCX.
"""
from __future__ import annotations
import datetime
from pathlib import Path

from parser import Document, TextRun


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------

def to_txt(doc: Document) -> str:
    lines = []
    if doc.header and doc.header.plain_text().strip():
        lines.append(doc.header.plain_text().strip(' \n'))
        lines.append('---')
    for para in doc.paragraphs:
        text = para.plain_text().strip(' \n')
        if text.strip():
            if para.page_break_before:
                lines.append('--- page break ---')
            if para.space_before and lines:
                lines.append('')
            lines.append(text)
    if doc.footer and doc.footer.plain_text().strip():
        lines.append('---')
        lines.append(doc.footer.plain_text().strip(' \n'))
    return '\n\n'.join(lines)


def save_txt(doc: Document, dest: Path) -> None:
    dest.write_text(to_txt(doc), encoding='utf-8')


# ---------------------------------------------------------------------------
# RTF
# ---------------------------------------------------------------------------

def _rtf_font_family(name: str) -> str:
    """Return the RTF font-family tag for a given font name."""
    n = name.lower()
    if any(w in n for w in ('sans', 'swiss', 'gothic', 'sanserif')):
        return r'\fswiss'
    if any(w in n for w in ('courier', 'prestige', 'condensed', 'mono')):
        return r'\fmodern'
    return r'\froman'


def _rtf_escape(text: str) -> str:
    """Escape special RTF characters."""
    out = []
    for ch in text:
        if ch == '\\':
            out.append('\\\\')
        elif ch == '{':
            out.append('\\{')
        elif ch == '}':
            out.append('\\}')
        elif ch == '\t':
            out.append(r'\tab ')
        elif ord(ch) > 127:
            out.append(f'\\u{ord(ch)}?')
        else:
            out.append(ch)
    return ''.join(out)


def _rtf_run(run: 'TextRun', font_idx: int | None = None) -> str:
    """Wrap a TextRun's escaped text in RTF formatting codes.

    *font_idx* is the RTF font table index (``\\fN``) to apply to this run.
    ``None`` means inherit the document default (``\\deff0``).
    """
    if run.page_number:
        content = r'\chpgn'
        if run.font_size is not None:
            return '{' + rf'\fs{int(run.font_size * 2)} ' + content + '}'
        return content
    if not run.text.strip():
        return ''
    escaped = _rtf_escape(run.text)
    pre, post = [], []
    if run.bold:
        pre.append(r'\b');       post.insert(0, r'\b0')
    if run.italic:
        pre.append(r'\i');       post.insert(0, r'\i0')
    if run.underline:
        pre.append(r'\ul');      post.insert(0, r'\ulnone')
    if run.superscript:
        pre.append(r'\super');   post.insert(0, r'\nosupersub')
    elif run.subscript:
        pre.append(r'\sub');     post.insert(0, r'\nosupersub')
    prefix = ' '.join(pre) + ' ' if pre else ''
    suffix = ' ' + ' '.join(post) if post else ''
    content = prefix + escaped + suffix
    if run.font_size is not None or font_idx is not None:
        # Wrap in an RTF group so font/size changes are scoped to this run only.
        size_code = rf'\fs{int(run.font_size * 2)}' if run.font_size is not None else ''
        font_code = rf'\f{font_idx}' if font_idx is not None else ''
        return '{' + font_code + size_code + ' ' + content + '}'
    return content


def _rtf_para(para: 'Paragraph', font_map: 'dict[str, int] | None' = None) -> str:
    """Render a single Paragraph as an RTF paragraph string (no trailing newline).

    *font_map* maps font face names to RTF font table indices.  Runs whose
    ``font_face`` appears in *font_map* receive an explicit ``\\fN`` switch.
    """
    _rtf_align = {'centre': r'\qc', 'right': r'\qr', 'left': ''}

    def _run(run):
        idx = font_map.get(run.font_face) if (font_map and run.font_face) else None
        return _rtf_run(run, idx)

    parts = [_run(run) for run in para.runs]
    parts = [p for p in parts if p]
    if not parts:
        return ''
    indent = rf'\li{para.left_indent}' if para.left_indent else ''
    font_size = rf'\fs{int(para.font_size * 2)}' if para.font_size is not None else ''
    page_break = r'\page' if para.page_break_before else ''
    if para.footer_tab:
        zone1 = [_run(r) for r in para.runs if r.page_number]
        zone2 = [_run(r) for r in para.runs if not r.page_number and r.text.strip()]
        content = r'\tab ' + ' '.join(zone1) + r' \tab ' + ' '.join(zone2)
        return page_break + r'\pard\tqc\tx4513\tqr\tx9026' + font_size + ' ' + content + r'\par'
    if para.inline_right_tab:
        return page_break + r'\pard\tqr\tx9026' + indent + font_size + ' ' + ' '.join(parts) + r'\par'
    align = _rtf_align.get(para.alignment, '')
    tab_stops = ''.join(rf'\tx{twips}' for twips in sorted(set(para.tab_stops)))
    return page_break + r'\pard' + align + indent + tab_stops + font_size + ' ' + ' '.join(parts) + r'\par'


def to_rtf(doc: Document) -> str:
    # Build font table from doc.fonts (slot 0 = \f0 default; slot 2 = \f1 alternate).
    default_font = doc.fonts[0] if doc.fonts and doc.fonts[0] else 'Times New Roman'
    alt_font = doc.fonts[2] if doc.fonts and len(doc.fonts) > 2 and doc.fonts[2] else ''

    font_entries = [
        rf'{{\f0{_rtf_font_family(default_font)}\fcharset0 {default_font};}}'
    ]
    font_map: dict[str, int] = {}
    if alt_font:
        font_entries.append(rf'{{\f1{_rtf_font_family(alt_font)}\fcharset0 {alt_font};}}')
        font_map[alt_font] = 1

    rtf_header = (
        r'{\rtf1\ansi\deff0'
        r'\paperw11906\paperh16838'
        r'\margl1440\margr1440\margt1440\margb1440'
        '{\\fonttbl' + ''.join(font_entries) + '}'
        r'{\colortbl ;}'
        '\n'
    )

    parts = []

    if doc.header and doc.header.plain_text().strip():
        hdr = _rtf_para(doc.header, font_map)
        if hdr:
            parts.append(r'{\header ' + hdr + '}\n')

    if doc.footer and doc.footer.plain_text().strip():
        ftr = _rtf_para(doc.footer, font_map)
        if ftr:
            parts.append(r'{\footer ' + ftr + '}\n')

    for para in doc.paragraphs:
        if not para.plain_text().strip():
            continue
        if para.space_before:
            parts.append(r'\pard\par' + '\n')
        p = _rtf_para(para, font_map)
        if p:
            parts.append(p + '\n')

    return rtf_header + ''.join(parts) + '}'


def save_rtf(doc: Document, dest: Path) -> None:
    dest.write_text(to_rtf(doc), encoding='ascii', errors='replace')


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def save_docx(doc: Document, dest: Path) -> None:
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        raise RuntimeError(
            "python-docx is required for DOCX output. "
            "Install it with: pip install python-docx"
        ) from e

    docx = DocxDocument()

    normal = docx.styles['Normal']
    if doc.fonts and doc.fonts[0]:
        normal.font.name = doc.fonts[0]
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_after = Pt(0)

    _docx_align = {
        'centre': WD_ALIGN_PARAGRAPH.CENTER,
        'right':  WD_ALIGN_PARAGRAPH.RIGHT,
    }

    def _apply_tab_stops(paragraph, tab_stops: list) -> None:
        """Add explicit left tab stops (in twips) to a paragraph's XML."""
        if not tab_stops:
            return
        pPr = paragraph._p.get_or_add_pPr()
        tabs_el = OxmlElement('w:tabs')
        for twips in sorted(set(tab_stops)):
            tab = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'left')
            tab.set(qn('w:pos'), str(twips))
            tabs_el.append(tab)
        pPr.append(tabs_el)

    def _apply_footer_tab_stops(paragraph) -> None:
        """Add centre+right tab stops for two-zone footer layout (hardcoded A4 1-inch margins)."""
        pPr = paragraph._p.get_or_add_pPr()
        tabs_el = OxmlElement('w:tabs')
        centre_tab = OxmlElement('w:tab')
        centre_tab.set(qn('w:val'), 'center')
        centre_tab.set(qn('w:pos'), '4513')
        tabs_el.append(centre_tab)
        right_tab = OxmlElement('w:tab')
        right_tab.set(qn('w:val'), 'right')
        right_tab.set(qn('w:pos'), '9026')
        tabs_el.append(right_tab)
        pPr.append(tabs_el)

    def _add_tab_run(p) -> None:
        """Append a bare tab character run to paragraph p."""
        r_el = OxmlElement('w:r')
        tab_el = OxmlElement('w:tab')
        r_el.append(tab_el)
        p._p.append(r_el)

    def _add_page_number_field(p, run):
        """Insert a PAGE field into paragraph p as three fldChar runs."""
        from lxml import etree
        for fld_type in ('begin', 'end'):
            r_el = OxmlElement('w:r')
            if run.font_size is not None:
                rpr = OxmlElement('w:rPr')
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), str(int(run.font_size * 2)))
                rpr.append(sz)
                r_el.append(rpr)
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), fld_type)
            r_el.append(fld)
            if fld_type == 'begin':
                # Insert instrText run between begin and end
                p._p.append(r_el)
                instr_el = OxmlElement('w:r')
                instr = OxmlElement('w:instrText')
                instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                instr.text = ' PAGE '
                instr_el.append(instr)
                p._p.append(instr_el)
            else:
                p._p.append(r_el)

    def _add_para(container, para):
        p = container.add_paragraph()
        if para.page_break_before:
            p.paragraph_format.page_break_before = True
        if para.footer_tab:
            _apply_footer_tab_stops(p)
            # Zone 1: leading tab + page number centred; Zone 2: tab + reference right-aligned.
            _add_tab_run(p)
            for run in para.runs:
                if run.page_number:
                    _add_page_number_field(p, run)
                elif run.text.strip():
                    _add_tab_run(p)
                    r = p.add_run(run.text)
                    effective_size = run.font_size if run.font_size is not None else para.font_size
                    if effective_size is not None:
                        r.font.size = Pt(effective_size)
            return p
        else:
            if para.alignment in _docx_align:
                p.alignment = _docx_align[para.alignment]
            if para.left_indent:
                p.paragraph_format.left_indent = Twips(para.left_indent)
            if para.inline_right_tab:
                # Right tab stop at right margin (hardcoded A4 1-inch margins: 9026 twips).
                pPr = p._p.get_or_add_pPr()
                tabs_el = OxmlElement('w:tabs')
                right_tab = OxmlElement('w:tab')
                right_tab.set(qn('w:val'), 'right')
                right_tab.set(qn('w:pos'), '9026')
                tabs_el.append(right_tab)
                pPr.append(tabs_el)
            else:
                _apply_tab_stops(p, para.tab_stops)
        for run in para.runs:
            if run.page_number:
                _add_page_number_field(p, run)
                continue
            if not run.text.strip():
                continue
            r = p.add_run(run.text)
            if run.font_face:    r.font.name = run.font_face
            if run.bold:         r.bold = True
            if run.italic:       r.italic = True
            if run.underline:    r.underline = True
            if run.superscript:  r.font.superscript = True
            if run.subscript:    r.font.subscript = True
            effective_size = run.font_size if run.font_size is not None else para.font_size
            if effective_size is not None:
                r.font.size = Pt(effective_size)

    section = docx.sections[0]

    if doc.header and doc.header.plain_text().strip():
        section.header.is_linked_to_previous = False
        # Clear the default empty paragraph python-docx adds
        for p in section.header.paragraphs:
            p.clear()
        _add_para(section.header, doc.header)

    if doc.footer and doc.footer.plain_text().strip():
        section.footer.is_linked_to_previous = False
        for p in section.footer.paragraphs:
            p.clear()
        _add_para(section.footer, doc.footer)

    for para in doc.paragraphs:
        if not para.plain_text().strip():
            continue
        if para.space_before:
            docx.add_paragraph()
        _add_para(docx, para)

    docx.save(dest)


# ---------------------------------------------------------------------------
# Dispatch
# ---------------------------------------------------------------------------

SUPPORTED_FORMATS = {
    '.txt':  save_txt,
    '.rtf':  save_rtf,
    '.docx': save_docx,
}


def convert(doc: Document, dest: Path) -> None:
    """Convert a parsed Document to the format implied by dest's extension."""
    ext = dest.suffix.lower()
    if ext not in SUPPORTED_FORMATS:
        raise ValueError(f"Unsupported output format: {ext!r}. Choose from {list(SUPPORTED_FORMATS)}")
    SUPPORTED_FORMATS[ext](doc, dest)


def log_error(log_path: Path, filename: str, error: Exception) -> None:
    """Append a timestamped error entry to the error log."""
    timestamp = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    entry = f"{timestamp} - {filename}: {type(error).__name__}: {error}\n"
    with open(log_path, 'a', encoding='utf-8') as f:
        f.write(entry)


def log_warning(log_path: Path, filename: str, message: str) -> None:
    """Append a timestamped warning entry to the error log."""
    timestamp = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    entry = f"{timestamp} - {filename}: [WARNING] {message}\n"
    with open(log_path, 'a', encoding='utf-8') as f:
        f.write(entry)
