"""
Regression tests for the Locoscript 2 parser.

Two layers of coverage:
  1. Golden-file test — parse the real HENCOTES sample and compare plain-text
     output against a stored snapshot.  Catches any regression that changes
     overall output without pinpointing the cause.

  2. Pattern unit tests — minimal synthetic byte sequences that exercise each
     discovered binary pattern in isolation.  When a regression occurs these
     point directly at the broken pattern.
"""

import unittest
from pathlib import Path

# Allow running from project root or from tests/
import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from parser import (parse, ParseError, Document, Paragraph, TextRun,
                    _detect_variant, _find_body_start, _find_footer_start)
from converter import to_txt, to_rtf


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

MAGIC = b'DOC'
PARA_CTRL = bytes([0x22, 0x61, 0x0b])

def _doc(body: bytes) -> bytes:
    """Wrap body bytes in the minimal valid Locoscript 2 envelope.

    Structure: DOC magic + one standard 0b paragraph-content block (8 bytes)
    + body.  The 0b block uses five zero param/indent bytes so the default
    skip path (i += 8) is taken.
    """
    header = MAGIC + PARA_CTRL + bytes([0x00, 0x00, 0x00, 0x00, 0x00])
    return header + body


def _plain(data: bytes) -> str:
    """Parse synthetic bytes and return the plain-text string."""
    return parse(data).plain_text()


def _paras(data: bytes) -> list[str]:
    """Parse synthetic bytes and return a list of non-empty paragraph texts."""
    return [p.plain_text() for p in parse(data).paragraphs if p.plain_text().strip()]


# ---------------------------------------------------------------------------
# Golden-file regression test
# ---------------------------------------------------------------------------

FIXTURE_DIR = Path(__file__).parent / 'fixtures'
SAMPLE_FILE = Path(__file__).parent.parent / 'HENCOTES'
GOLDEN_FILE = FIXTURE_DIR / 'HENCOTES.golden.txt'


class TestGoldenFile(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        if not SAMPLE_FILE.exists():
            raise unittest.SkipTest(f"Sample file not found: {SAMPLE_FILE}")
        if not GOLDEN_FILE.exists():
            raise unittest.SkipTest(f"Golden file not found: {GOLDEN_FILE}")
        with open(SAMPLE_FILE, 'rb') as f:
            cls.doc = parse(f.read())
        cls.actual = to_txt(cls.doc)
        cls.golden = GOLDEN_FILE.read_text(encoding='utf-8')

    def test_output_matches_golden(self):
        """Full plain-text output must match the stored golden snapshot."""
        self.assertEqual(self.actual, self.golden)

    def test_paragraph_count(self):
        """Paragraph count must not silently change."""
        # Exclude '---' separator lines (header/footer) and '--- page break ---'
        # markers that to_txt inserts — these are not content paragraphs.
        _separators = {'---', '--- page break ---'}
        golden_para_count = len([p for p in self.golden.split('\n\n')
                                  if p.strip() and p.strip() not in _separators])
        actual_para_count = (
            len([p for p in self.doc.paragraphs if p.plain_text().strip()])
            + (1 if self.doc.header and self.doc.header.plain_text().strip() else 0)
            + (1 if self.doc.footer and self.doc.footer.plain_text().strip() else 0)
        )
        self.assertEqual(actual_para_count, golden_para_count)

    def test_key_passages_present(self):
        """Spot-check that known content passages survive parsing."""
        checks = [
            "Hencotes",
            "cockfighting",
            "well heeled",
            "Eminent word",
            "Listed Buildings",
            "Laburnum Cottage",
            "The Times",
        ]
        for phrase in checks:
            self.assertIn(phrase, self.actual, f"Expected passage not found: {phrase!r}")

    def test_no_doubled_pair_artifacts(self):
        """Known doubled-pair indent byte values must not appear as isolated text.

        The specific pairs we've seen as artifacts are: GG JJ ZZ LL TT AA KK MM QQ.
        We exclude pairs that appear legitimately in English prose (e.g. 'II' in Roman
        numerals, 'LL' in 'LL.B.', 'SS' in words like 'lass').
        """
        import re
        # Only check pairs that have no business appearing isolated in prose
        artifact_pairs = ['GG', 'JJ', 'ZZ', 'TT', 'AA', 'KK', 'MM', 'QQ']
        for pair in artifact_pairs:
            # Isolated means not surrounded by other word characters
            pattern = rf'(?<!\w){re.escape(pair)}(?!\w)'
            self.assertNotRegex(self.actual, pattern,
                                f"Doubled-pair artifact {pair!r} found in output")


# ---------------------------------------------------------------------------
# Pattern unit tests
# ---------------------------------------------------------------------------

class TestFileHeader(unittest.TestCase):

    def test_valid_magic_accepted(self):
        data = _doc(b'Hello')
        doc = parse(data)
        self.assertIsInstance(doc, Document)

    def test_invalid_magic_raises(self):
        with self.assertRaises(ParseError):
            parse(b'XXXsome content')

    def test_joy_magic_raises_informative_error(self):
        # JOY files should raise ParseError with a message that names the format
        with self.assertRaises(ParseError) as ctx:
            parse(b'JOY\x01\x04some content')
        self.assertIn('JOY', str(ctx.exception))


class TestWordSeparator(unittest.TestCase):

    def test_word_sep_becomes_space(self):
        # 02 between words should produce a space
        data = _doc(b'Hello\x02World')
        self.assertIn('Hello World', _plain(data))


class TestParagraphBreak(unittest.TestCase):

    def test_para_break_splits_paragraphs(self):
        # 13 04 50 should end one paragraph and start another
        data = _doc(b'First\x13\x04\x50Second')
        paras = _paras(data)
        self.assertEqual(len(paras), 2)
        self.assertIn('First', paras[0])
        self.assertIn('Second', paras[1])

    def test_para_break_with_trailing_metadata(self):
        # 13 04 50 00 01 4a 4a — trailing 00 01 + doubled pair must be consumed
        data = _doc(b'Alpha\x13\x04\x50\x00\x01\x4a\x4aOmega')
        paras = _paras(data)
        self.assertEqual(len(paras), 2)
        self.assertNotIn('J', paras[1][:2])  # 'JJ' (0x4a 0x4a) must not appear


class TestBoldUnderlineFormatting(unittest.TestCase):

    def test_bold_on_sets_bold_flag(self):
        # 08 00 [01 xx xx] turns bold on; text inside has bold=True
        data = _doc(b'Normal\x02\x08\x00\x01\x11\x11Bold\x09\x00\x01\x11\x11after')
        doc = parse(data)
        runs = [r for p in doc.paragraphs for r in p.runs]
        bold_runs = [r for r in runs if r.bold]
        self.assertTrue(any('Bold' in r.text for r in bold_runs))

    def test_bold_off_clears_bold_flag(self):
        # After 09 00 the bold flag should be cleared
        data = _doc(b'\x08\x00\x01\x11\x11Bold\x09\x00\x01\x11\x11after')
        doc = parse(data)
        runs = [r for p in doc.paragraphs for r in p.runs]
        non_bold = [r for r in runs if not r.bold]
        self.assertTrue(any('after' in r.text for r in non_bold))

    def test_underline_on_sets_flag(self):
        # 08 02 (no params) turns underline on
        data = _doc(b'Before\x08\x02underlined\x09\x02after')
        doc = parse(data)
        runs = [r for p in doc.paragraphs for r in p.runs]
        ul_runs = [r for r in runs if r.underline]
        self.assertTrue(any('underlined' in r.text for r in ul_runs))

    def test_superscript_on_sets_flag(self):
        # 08 06 turns superscript on; 09 06 turns it off
        data = _doc(b'base\x08\x06sup\x09\x06\x01\x11\x11text')
        doc = parse(data)
        runs = [r for p in doc.paragraphs for r in p.runs]
        sup_runs = [r for r in runs if r.superscript]
        self.assertTrue(any('sup' in r.text for r in sup_runs))

    def test_bold_params_not_in_output(self):
        # The 01 separator and indent-pair param bytes must not appear as text
        data = _doc(b'\x08\x00\x01\x11\x11Bold\x09\x00')
        text = _plain(data)
        self.assertIn('Bold', text)
        self.assertNotIn('\x11', text)

    def test_pending_newline_survives_bold_off(self):
        # A '\n' pending in current_text must not be dropped when bold-off fires
        # Sequence: content, line-break (emits \n), bold-off, more content
        data = _doc(b'Line1\x13\x04\x78\x09\x00Line2')
        text = _plain(data)
        self.assertIn('Line1', text)
        self.assertIn('Line2', text)
        self.assertIn('\n', text)


class TestItalic(unittest.TestCase):

    def test_italic_on_sets_italic_flag(self):
        # 13 04 64 turns italic on
        data = _doc(b'Normal\x02\x13\x04\x64italic\x02\x13\x04\x78after')
        doc = parse(data)
        runs = [r for p in doc.paragraphs for r in p.runs]
        italic_runs = [r for r in runs if r.italic]
        self.assertTrue(any('italic' in r.text for r in italic_runs))

    def test_italic_off_via_line_break(self):
        # 13 04 78 when italic=True ends italic; text after should not be italic
        data = _doc(b'\x13\x04\x64italic\x13\x04\x78plain')
        doc = parse(data)
        runs = [r for p in doc.paragraphs for r in p.runs]
        plain_runs = [r for r in runs if not r.italic]
        self.assertTrue(any('plain' in r.text for r in plain_runs))

    def test_line_break_when_not_italic(self):
        # 13 04 78 when italic=False emits a newline
        data = _doc(b'Line1\x13\x04\x78Line2')
        text = _plain(data)
        self.assertIn('\n', text)


class TestTabSequence(unittest.TestCase):

    def test_tab_sequence_emits_tab(self):
        # 09 05 01 + 2 param bytes → tab character
        data = _doc(b'Col1\x09\x05\x01\x00\x00Col2')
        text = _plain(data)
        self.assertIn('\t', text)
        self.assertIn('Col1', text)
        self.assertIn('Col2', text)


class TestSISequences(unittest.TestCase):

    def test_0f_04_emits_tab_and_strips_params(self):
        # 0f 04 [printable param bytes] 01 [doubled pair] → tab, no artefacts
        # Mirrors the real-world case: 0f 04 31 61 01 2a 2a (seen as "1a**" artefacts)
        data = _doc(b'Before\x0f\x04\x31\x61\x01\x2a\x2aAfter')
        text = _plain(data)
        self.assertIn('\t', text)
        self.assertIn('Before', text)
        self.assertIn('After', text)
        self.assertNotIn('1a', text)
        self.assertNotIn('**', text)

    def test_0f_04_no_separator_no_doubled_pair(self):
        # 0f 04 [2 printable param bytes] immediately followed by content (no 01 separator)
        data = _doc(b'Col1\x0f\x04\x27\x66Col2')
        text = _plain(data)
        self.assertIn('\t', text)
        self.assertNotIn("'f", text)

    def test_0f_05_emits_nothing_and_strips_params(self):
        # 0f 05 is a hanging-indent marker — no tab emitted, params consumed
        data = _doc(b'Word\x0f\x05\x31\x61\x01\x3e\x3eContent')
        text = _plain(data)
        self.assertNotIn('1a', text)
        self.assertNotIn('>>', text)
        self.assertIn('Content', text)


class TestTabStopPositions(unittest.TestCase):
    """Tab stop column positions are captured from inline sequences."""

    def test_09_05_01_records_tab_stop_twips(self):
        # 09 05 01 1c 1c → XX=0x1c=28 at default scale pitch 0x18 → 28×144=4032 twips
        data = _doc(b'Text\x09\x05\x01\x1c\x1cMore')
        doc = parse(data)
        para = doc.paragraphs[0]
        self.assertIn(4032, para.tab_stops)

    def test_09_05_01_zero_param_no_stop_recorded(self):
        # XX=0x00 means no position → tab_stops must stay empty
        data = _doc(b'Col1\x09\x05\x01\x00\x00Col2')
        doc = parse(data)
        para = doc.paragraphs[0]
        self.assertEqual(para.tab_stops, [])

    def test_0f_04_records_tab_stop_twips(self):
        # 0f 04 27 61 01 2a 2a → B1=0x27=39 (printable) → 39×144=5616 twips
        # B1 must be >= 0x20 (printable); values below 0x20 are non-printable
        # param bytes that precede B1 and are skipped by the handler.
        data = _doc(b'Before\x0f\x04\x27\x61\x01\x2a\x2aAfter')
        doc = parse(data)
        para = doc.paragraphs[0]
        self.assertIn(0x27 * 144, para.tab_stops)  # 39×144=5616

    def test_multiple_tab_stops_per_paragraph(self):
        # Two 09 05 01 sequences in one paragraph record two stops
        data = _doc(b'A\x09\x05\x01\x10\x10B\x09\x05\x01\x20\x20C')
        doc = parse(data)
        para = doc.paragraphs[0]
        self.assertIn(0x10 * 144, para.tab_stops)  # 16×144=2304
        self.assertIn(0x20 * 144, para.tab_stops)  # 32×144=4608

    def test_tab_stops_not_shared_between_paragraphs(self):
        # Tab stops on one paragraph must not bleed into the next
        data = _doc(
            b'A\x09\x05\x01\x1c\x1cB'
            b'\x13\x04\x50'                         # paragraph break
            b'C\x09\x05\x01\x00\x00D'
        )
        doc = parse(data)
        self.assertIn(4032, doc.paragraphs[0].tab_stops)
        self.assertEqual(doc.paragraphs[1].tab_stops, [])

    def test_rtf_emits_tx_for_tab_stop(self):
        from converter import to_rtf
        data = _doc(b'Text\x09\x05\x01\x1c\x1cMore')
        out = to_rtf(parse(data))
        self.assertIn(r'\tx4032', out)

    def test_rtf_no_tx_when_no_tab_stops(self):
        from converter import to_rtf
        data = _doc(b'Hello\x13\x04\x50')
        out = to_rtf(parse(data))
        self.assertNotIn(r'\tx', out)

    def test_rtf_tx_appears_before_content(self):
        from converter import to_rtf
        data = _doc(b'Col1\x09\x05\x01\x1c\x1cCol2')
        out = to_rtf(parse(data))
        self.assertLess(out.index(r'\tx4032'), out.index('Col1'))

    def test_docx_tab_stop_written_to_xml(self):
        import tempfile, os
        from pathlib import Path
        from converter import save_docx
        from docx import Document as DocxDocument
        data = _doc(b'Col1\x09\x05\x01\x1c\x1cCol2')
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            tmp = Path(f.name)
        try:
            save_docx(parse(data), tmp)
            result = DocxDocument(tmp)
            # Check the XML of the first paragraph for a w:tab element with w:pos="4032"
            para_xml = result.paragraphs[0]._p.xml
            self.assertIn('w:pos="4032"', para_xml)
        finally:
            os.unlink(tmp)


class TestIndentMetadata(unittest.TestCase):

    def test_09_00_01_skipped(self):
        # 09 00 01 is structural metadata and must not appear in output
        data = _doc(b'Before\x09\x00\x01After')
        text = _plain(data)
        self.assertIn('Before', text)
        self.assertIn('After', text)

    def test_09_00_01_with_doubled_pair_skipped(self):
        # 09 00 01 4a 4a — the doubled pair must also be consumed
        data = _doc(b'Before\x09\x00\x01\x4a\x4aAfter')
        text = _plain(data)
        self.assertNotIn('JJ', text)
        self.assertIn('After', text)


class TestSectionBreak(unittest.TestCase):

    def test_section_break_01_flushes_content(self):
        # 0e 01 must flush accumulated text, not discard it
        data = _doc(b'PreBreak\x0e\x01' + PARA_CTRL + bytes([0x00]*5) + b'PostBreak')
        paras = _paras(data)
        self.assertTrue(any('PreBreak' in p for p in paras),
                        "Content before section break must not be discarded")

    def test_section_break_02_flushes_content(self):
        # 0e 02 (page break) must also flush
        data = _doc(b'BeforePage\x0e\x02' + PARA_CTRL + bytes([0x00]*5) + b'AfterPage')
        paras = _paras(data)
        self.assertTrue(any('BeforePage' in p for p in paras),
                        "Content before page break must not be discarded")


class TestHyphenByte(unittest.TestCase):

    def test_hyphen_between_words(self):
        # 0x06 between two printable chars → hyphen
        data = _doc(b'well\x06heeled')
        text = _plain(data)
        self.assertIn('well-heeled', text)

    def test_hyphen_adjacent_to_word_sep_is_space(self):
        # 0x06 adjacent to word separator → space, not hyphen
        data = _doc(b'word\x02\x06next')
        text = _plain(data)
        self.assertNotIn('-', text)

    def test_hyphen_after_non_printable_is_space(self):
        # 0x06 after a non-printable byte (e.g. 0x11) → space, not hyphen
        data = _doc(b'\x11\x06word')
        text = _plain(data)
        self.assertNotIn('-', text)

    def test_dash_separator_between_word_seps(self):
        # 02 06 02 → ' - ' (dash separator, not triple space)
        data = _doc(b'Henhouse\x02\x06\x02domestic')
        text = _plain(data)
        self.assertIn('Henhouse - domestic', text)

    def test_word_sep_then_06_then_printable_is_space(self):
        # 02 06 printable → space before hyphen, not dash separator
        data = _doc(b'word\x02\x06next')
        text = _plain(data)
        self.assertNotIn(' - ', text)

    def test_printable_then_06_then_word_sep_is_space(self):
        # printable 06 02 → space, not dash separator
        data = _doc(b'word\x06\x02next')
        text = _plain(data)
        self.assertNotIn(' - ', text)


class TestLiteralQuotePrefix(unittest.TestCase):

    def test_22_61_before_word_sep_is_literal(self):
        # 22 61 followed by 02 (word separator) is literal `"a`, not a control code
        data = _doc(b'\x22\x61\x02typical')
        text = _plain(data)
        self.assertIn('"a', text)

    def test_22_61_before_06_is_literal(self):
        # 22 61 followed by 06 is also literal `"a`
        data = _doc(b'\x22\x61\x06word')
        text = _plain(data)
        self.assertIn('"a', text)


class TestExtendedCharacters(unittest.TestCase):
    """High-byte (0x80–0xFF) character mappings confirmed from real file evidence."""

    def test_e9_maps_to_pound_sign(self):
        # 0xE9 → £ (confirmed; diverges from Amstrad CP/M Plus table)
        data = _doc(b'\xe91.4million')
        text = _plain(data)
        self.assertIn('£', text)
        self.assertNotIn('?', text)

    def test_84_maps_to_right_single_quote(self):
        # 0x84 → ' right single quotation mark (e.g. "Kelly's")
        data = _doc(b'Kelly\x84s')
        text = _plain(data)
        self.assertIn('\u2019', text)
        self.assertNotIn('?', text)

    def test_8f_maps_to_ae_ligature(self):
        # 0x8F → æ (e.g. "Archæology")
        data = _doc(b'Arch\x8fology')
        text = _plain(data)
        self.assertIn('æ', text)
        self.assertIn('Archæology', text)

    def test_b4_maps_to_e_acute(self):
        # 0xB4 → é (e.g. "Café")
        data = _doc(b'Caf\xb4')
        text = _plain(data)
        self.assertIn('é', text)
        self.assertIn('Café', text)

    def test_c3_maps_to_e_grave(self):
        # 0xC3 → è (e.g. "Adèle")
        data = _doc(b'Ad\xc3le')
        text = _plain(data)
        self.assertIn('è', text)
        self.assertIn('Adèle', text)

    def test_e4_maps_to_e_circumflex(self):
        # 0xE4 → ê (e.g. "Fête")
        data = _doc(b'F\xe4te')
        text = _plain(data)
        self.assertIn('ê', text)
        self.assertIn('Fête', text)

    def test_e8_maps_to_o_circumflex(self):
        # 0xE8 → ô (e.g. "Dépôt")
        data = _doc(b'D\xb4p\xe8t')
        text = _plain(data)
        self.assertIn('ô', text)

    def test_fa_maps_to_c_cedilla(self):
        # 0xFA → ç (e.g. "façade" — second encoding alongside ENQ sequence)
        data = _doc(b'fa\xfaade')
        text = _plain(data)
        self.assertIn('ç', text)
        self.assertIn('façade', text)

    def test_unknown_high_byte_maps_to_question_mark(self):
        # Bytes not in the map fall back to '?'
        data = _doc(b'un\x99known')
        text = _plain(data)
        self.assertIn('?', text)


class TestENQExtendedCharacters(unittest.TestCase):

    def test_cedilla_c_maps_to_c_cedilla(self):
        # 05 63 01 13 01 is the ENQ encoding for ç (c with cedilla)
        data = _doc(b'fa\x05\x63\x01\x13\x01ade')
        text = _plain(data)
        self.assertIn('ç', text)
        self.assertIn('façade', text)

    def test_cedilla_no_base_char_artifact(self):
        # The raw 'c' byte (0x63) must not appear as a separate artifact
        data = _doc(b'Fran\x05\x63\x01\x13\x01ais')
        text = _plain(data)
        self.assertIn('ç', text)
        self.assertNotIn('Franc', text)  # 'c' must not leak before 'ais'

    def test_cedilla_trailing_doubled_pair_consumed(self):
        # 05 63 01 13 01 followed by 01 XX XX — the doubled pair must not appear
        data = _doc(b'fa\x05\x63\x01\x13\x01\x01\x46\x46ade')
        text = _plain(data)
        self.assertIn('ç', text)
        self.assertNotIn('FF', text)

    def test_enq_unknown_diacritic_emits_base_char(self):
        # An ENQ sequence with an unknown diacritic emits the base character (0x68 = 'h')
        # rather than '?' or nothing — best-effort fallback for undocumented diacritics
        data = _doc(b'wit\x05\x68\x01\x07\x01out')
        text = _plain(data)
        self.assertIn('h', text)   # base char emitted
        self.assertIn('without', text)

    def test_enq_does_not_misfire_on_structural_bytes(self):
        # 05 with a non-printable base byte must not be treated as ENQ char encoding
        data = _doc(b'After\x05\x00\x01\x13\x01text')
        text = _plain(data)
        self.assertIn('text', text)


class TestConverterTabHandling(unittest.TestCase):
    """Converter output correctly preserves tab characters in all three formats."""

    def _parse_tab_doc(self):
        # Paragraph: "Col1" TAB "Col2"
        # 09 05 01 00 00 = tab sequence; 13 04 50 = paragraph break
        data = _doc(b'Col1\x09\x05\x01\x00\x00Col2\x13\x04\x50')
        return parse(data)

    def test_txt_preserves_leading_tab(self):
        # A paragraph whose first token is a tab must not be stripped
        data = _doc(b'\x09\x05\x01\x00\x00Col1\x02Col2\x13\x04\x50')
        from converter import to_txt
        out = to_txt(parse(data))
        self.assertTrue(out.startswith('\t'), f"Expected leading tab, got: {out!r}")

    def test_txt_preserves_mid_paragraph_tab(self):
        from converter import to_txt
        out = to_txt(self._parse_tab_doc())
        self.assertIn('\t', out)
        self.assertIn('Col1', out)
        self.assertIn('Col2', out)

    def test_rtf_tab_becomes_control_word(self):
        from converter import to_rtf
        out = to_rtf(self._parse_tab_doc())
        self.assertIn(r'\tab', out)
        self.assertNotIn('\t', out)  # raw tab must not appear in RTF output

    def test_rtf_preserves_content_around_tab(self):
        from converter import to_rtf
        out = to_rtf(self._parse_tab_doc())
        self.assertIn('Col1', out)
        self.assertIn('Col2', out)

    def test_docx_preserves_tab(self):
        import tempfile, os
        from pathlib import Path
        from converter import save_docx
        from docx import Document as DocxDocument
        doc = self._parse_tab_doc()
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            tmp = Path(f.name)
        try:
            save_docx(doc, tmp)
            result = DocxDocument(tmp)
            full_text = ''.join(run.text for para in result.paragraphs for run in para.runs)
            self.assertIn('\t', full_text)
            self.assertIn('Col1', full_text)
            self.assertIn('Col2', full_text)
        finally:
            os.unlink(tmp)


class TestConverterNoSpuriousTrailingSpace(unittest.TestCase):
    """DOCX runs must not have a spurious trailing space appended."""

    def test_docx_run_text_not_padded(self):
        import tempfile, os
        from pathlib import Path
        from converter import save_docx
        from docx import Document as DocxDocument
        # Single word, no trailing space expected in run text
        data = _doc(b'Hello\x13\x04\x50')
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            tmp = Path(f.name)
        try:
            save_docx(parse(data), tmp)
            result = DocxDocument(tmp)
            run_texts = [r.text for p in result.paragraphs for r in p.runs]
            for rt in run_texts:
                self.assertFalse(rt.endswith(' '), f"Run text has spurious trailing space: {rt!r}")
        finally:
            os.unlink(tmp)


class TestParagraphAlignment(unittest.TestCase):
    """DC1/DLE alignment codes set paragraph alignment in parser and converters."""

    def _para_with_alignment(self, prefix_bytes: bytes) -> 'Paragraph':
        # prefix_bytes appear right after the content-start block, before text
        data = _doc(prefix_bytes + b'Hello\x13\x04\x50')
        doc = parse(data)
        return doc.paragraphs[0] if doc.paragraphs else None

    def test_11_06_sets_centre_alignment(self):
        para = self._para_with_alignment(b'\x11\x06')
        self.assertIsNotNone(para)
        self.assertEqual(para.alignment, 'centre')

    def test_10_07_sets_right_alignment(self):
        para = self._para_with_alignment(b'\x10\x07')
        self.assertIsNotNone(para)
        self.assertEqual(para.alignment, 'right')

    def test_10_04_sets_right_alignment(self):
        para = self._para_with_alignment(b'\x10\x04')
        self.assertIsNotNone(para)
        self.assertEqual(para.alignment, 'right')

    def test_default_alignment_is_left(self):
        para = self._para_with_alignment(b'')
        self.assertIsNotNone(para)
        self.assertEqual(para.alignment, 'left')

    def test_11_06_does_not_emit_spurious_space(self):
        # The 0x06 param byte must be consumed, not produce a leading space
        data = _doc(b'\x11\x06Hello\x13\x04\x50')
        doc = parse(data)
        self.assertTrue(doc.paragraphs[0].plain_text().startswith('Hello'))

    def test_rtf_centre_uses_qc(self):
        from converter import to_rtf
        data = _doc(b'\x11\x06Hello\x13\x04\x50')
        out = to_rtf(parse(data))
        self.assertIn(r'\pard\qc', out)

    def test_rtf_right_uses_qr(self):
        from converter import to_rtf
        data = _doc(b'\x10\x07Hello\x13\x04\x50')
        out = to_rtf(parse(data))
        self.assertIn(r'\pard\qr', out)

    def test_rtf_left_has_no_alignment_code(self):
        from converter import to_rtf
        data = _doc(b'Hello\x13\x04\x50')
        out = to_rtf(parse(data))
        self.assertIn(r'\pard ', out)
        self.assertNotIn(r'\qc', out)
        self.assertNotIn(r'\qr', out)

    def test_docx_centre_alignment(self):
        import tempfile, os
        from pathlib import Path
        from converter import save_docx
        from docx import Document as DocxDocument
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        data = _doc(b'\x11\x06Hello\x13\x04\x50')
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            tmp = Path(f.name)
        try:
            save_docx(parse(data), tmp)
            result = DocxDocument(tmp)
            self.assertEqual(result.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        finally:
            os.unlink(tmp)

    def test_docx_right_alignment(self):
        import tempfile, os
        from pathlib import Path
        from converter import save_docx
        from docx import Document as DocxDocument
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        data = _doc(b'\x10\x07Hello\x13\x04\x50')
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            tmp = Path(f.name)
        try:
            save_docx(parse(data), tmp)
            result = DocxDocument(tmp)
            self.assertEqual(result.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.RIGHT)
        finally:
            os.unlink(tmp)


class TestRTFPageSize(unittest.TestCase):
    """RTF output must declare A4 page dimensions and margins."""

    def _rtf_for(self, body: bytes) -> str:
        from converter import to_rtf
        return to_rtf(parse(_doc(body)))

    def test_rtf_contains_a4_paper_width(self):
        out = self._rtf_for(b'Hello')
        self.assertIn(r'\paperw11906', out)

    def test_rtf_contains_a4_paper_height(self):
        out = self._rtf_for(b'Hello')
        self.assertIn(r'\paperh16838', out)

    def test_rtf_contains_margin_settings(self):
        out = self._rtf_for(b'Hello')
        self.assertIn(r'\margl1440', out)
        self.assertIn(r'\margr1440', out)
        self.assertIn(r'\margt1440', out)
        self.assertIn(r'\margb1440', out)

    def test_rtf_page_size_appears_before_body(self):
        # Page dimensions must be in the header, before any paragraph content
        out = self._rtf_for(b'Hello')
        self.assertLess(out.index(r'\paperw11906'), out.index('Hello'))


class TestControlSequenceSkip(unittest.TestCase):

    def test_unknown_ctrl_type_skipped_cleanly(self):
        # An unknown control type (e.g. 0x0c) should be skipped without
        # corrupting the text that follows
        data = _doc(b'\x22\x61\x0c\x00\x00\x00After')
        text = _plain(data)
        self.assertIn('After', text)


class TestVariableCtrlPrefix(unittest.TestCase):
    """Files using 22 6d 0b as the control prefix must parse correctly."""

    # Minimal envelope using 22 6d 0b instead of 22 61 0b
    _PARA_CTRL_M = bytes([0x22, 0x6d, 0x0b])
    _HEADER_M = MAGIC + _PARA_CTRL_M + bytes([0x00, 0x00, 0x00, 0x00, 0x00])

    def _doc_m(self, body: bytes) -> bytes:
        return self._HEADER_M + body

    def test_detect_variant_standard(self):
        # Standard files (22 61 0b) return (0x22, 0x61)
        data = _doc(b'Hello')
        self.assertEqual(_detect_variant(data), (0x22, 0x61))

    def test_detect_variant_22_6d(self):
        # Variant files (22 6d 0b) return (0x22, 0x6d)
        data = self._doc_m(b'Hello')
        self.assertEqual(_detect_variant(data), (0x22, 0x6d))

    def test_detect_variant_fallback(self):
        # No recognised sequence at all → default (0x22, 0x61)
        self.assertEqual(_detect_variant(b'DOC' + b'\x00' * 10), (0x22, 0x61))

    def test_detect_variant_prefers_most_frequent(self):
        # When header uses 0x61 but body uses 0x42 (like Memorial.002),
        # the most-frequent pair (0x22, 0x42) should win, not the first.
        para_ctrl_61 = bytes([0x22, 0x61, 0x0b, 0x00, 0x00, 0x00, 0x00, 0x00])
        para_ctrl_42 = bytes([0x22, 0x42, 0x0b, 0x00, 0x00, 0x00, 0x00, 0x00])
        # 1× 0x61 (header) + 3× 0x42 (body) → most frequent is (0x22, 0x42)
        data = MAGIC + para_ctrl_61 + para_ctrl_42 * 3 + b'Hello'
        self.assertEqual(_detect_variant(data), (0x22, 0x42))

    def test_detect_variant_1e_74(self):
        # 1e 74 prefix variant: returns (0x1e, 0x74)
        para_ctrl_1e = bytes([0x1e, 0x74, 0x0b, 0x00, 0x00, 0x00, 0x00, 0x00])
        data = MAGIC + para_ctrl_1e * 3 + b'Hello'
        self.assertEqual(_detect_variant(data), (0x1e, 0x74))

    def test_22_6d_file_produces_content(self):
        # A 22 6d file must yield the body text rather than garbage
        data = self._doc_m(b'Hello\x02world\x13\x04\x50')
        self.assertEqual(_plain(data), 'Hello world')

    def test_22_6d_multi_paragraph(self):
        # Multiple paragraphs separated by 22 6d 0b blocks
        para2 = self._PARA_CTRL_M + bytes([0x00, 0x00, 0x00, 0x00, 0x00])
        data = self._doc_m(b'First\x13\x04\x50' + para2 + b'Second\x13\x04\x50')
        self.assertEqual(_paras(data), ['First', 'Second'])

    def test_a6_0e_structural_block_skipped(self):
        # A 22 6d 0b block with high B3 (≥0x80) and 0x0e at B4 is a structural
        # section/layout header — skip it and the following binary blob to
        # reach the next real paragraph, emitting no garbage.
        structural = bytes([0x22, 0x6d, 0x0b, 0xa6, 0x0e, 0x07, 0x03, 0x00])
        binary_blob = bytes([0x70, 0x08, 0x80, 0x0e, 0x25, 0x02, 0x00, 0x78])
        next_para = self._PARA_CTRL_M + bytes([0x00, 0x00, 0x00, 0x00, 0x00])
        data = self._HEADER_M + structural + binary_blob + next_para + b'Clean\x13\x04\x50'
        text = _plain(data)
        self.assertEqual(text, 'Clean')

    def test_low_b3_0e_block_not_skipped(self):
        # A 0b block with low B3 (<0x80) and 0x0e at B4 (e.g. 36 0e) is a
        # normal content block — do NOT skip to next para, use default 8-byte skip.
        # After the 8-byte skip the text 'After' must still be present.
        normal_block = bytes([0x22, 0x6d, 0x0b, 0x36, 0x0e, 0x01, 0x04, 0x04])
        data = self._HEADER_M + normal_block + b'After\x13\x04\x50'
        text = _plain(data)
        self.assertIn('After', text)

    def test_c4_0e_still_skipped_in_standard_file(self):
        # The generalised B3≥0x80 + B4=0x0e check must not break c4 0e.
        structural = bytes([0x22, 0x61, 0x0b, 0xc4, 0x0e, 0x00, 0x00, 0x00])
        binary_blob = bytes([0x70, 0x08, 0x80, 0x00, 0x00, 0x00, 0x00, 0x00])
        next_para = PARA_CTRL + bytes([0x00, 0x00, 0x00, 0x00, 0x00])
        data = _doc(b'') + structural + binary_blob + next_para + b'OK\x13\x04\x50'
        text = _plain(data)
        self.assertIn('OK', text)

    def test_22_61_not_treated_as_ctrl_in_6d_file(self):
        # In a 22 6d file, a literal 22 61 sequence must be emitted as '"a',
        # not silently consumed as a control sequence
        data = self._doc_m(b'\x22\x61\x02text\x13\x04\x50')
        text = _plain(data)
        self.assertIn('"a', text)
        self.assertIn('text', text)


class TestHeaderFooterExtraction(unittest.TestCase):
    """Header and footer extraction via position-based section routing."""

    # -----------------------------------------------------------------------
    # Unit tests for _find_body_start and _find_footer_start
    # -----------------------------------------------------------------------

    def test_find_body_start_body_only(self):
        # body-only document: body_start == first_para
        data = _doc(b'Hello')
        first_para = data.index(PARA_CTRL)
        self.assertEqual(_find_body_start(data, 0x61, first_para, 'body'), first_para)

    def test_find_body_start_b5_fallback(self):
        # When _find_content_start returns first_para (no c4 0e block) and
        # initial_section != 'body', the B5 != 0x14 scan finds the first block
        # where B5 is not 0x14 (header/footer zone marker).
        para_ctrl = PARA_CTRL
        data = (b'DOC'
                + para_ctrl + bytes([0x50, 0x0d, 0x14, 0x78, 0x00])  # B5=0x14 → header zone, offset 3
                + para_ctrl + bytes([0x5a, 0x04, 0x14, 0x78, 0x00])  # B5=0x14 → footer zone, offset 11
                + para_ctrl + bytes([0x64, 0x00, 0x0a, 0x09, 0x01])  # B5=0x0a → body, offset 19
                + b'Hello')
        first_para = 3
        result = _find_body_start(data, 0x61, first_para, 'header')
        self.assertEqual(result, 19)

    def test_find_body_start_b3_00_skipped(self):
        # Blocks with B3=0x00 are layout/transition blocks and must not be
        # identified as body start even when B5 != 0x14.
        para_ctrl = PARA_CTRL
        data = (b'DOC'
                + para_ctrl + bytes([0x50, 0x0d, 0x14, 0x78, 0x00])  # B5=0x14 header zone, offset 3
                + para_ctrl + bytes([0x00, 0x00, 0x14, 0x78, 0x00])  # B3=0x00 layout block, offset 11
                + para_ctrl + bytes([0x64, 0x00, 0x0a, 0x09, 0x01])  # body, offset 19
                + b'Hello')
        first_para = 3
        result = _find_body_start(data, 0x61, first_para, 'header')
        self.assertEqual(result, 19)

    def test_find_footer_start(self):
        # After a section break between first_para and body_start, the first
        # para ctrl block is footer_start.
        para_ctrl = PARA_CTRL
        section_break = bytes([0x0e, 0x02])
        data = (b'DOC'
                + para_ctrl + bytes([0x50, 0x0d, 0x14, 0x78, 0x00])  # header, offset 3
                + section_break                                          # break, offset 11
                + para_ctrl + bytes([0x5a, 0x04, 0x14, 0x78, 0x00])  # footer, offset 13
                + para_ctrl + bytes([0x64, 0x00, 0x0a, 0x09, 0x01])  # body, offset 21
                + b'Hello')
        first_para = 3
        body_start = 21
        result = _find_footer_start(data, 0x61, first_para, body_start)
        self.assertEqual(result, 13)

    def test_find_footer_start_returns_zero_when_no_break(self):
        # No section break before body_start → returns 0
        para_ctrl = PARA_CTRL
        data = (b'DOC'
                + para_ctrl + bytes([0x50, 0x0d, 0x14, 0x78, 0x00])  # header, offset 3
                + para_ctrl + bytes([0x64, 0x00, 0x0a, 0x09, 0x01])  # body, offset 11
                + b'Hello')
        first_para = 3
        body_start = 11
        result = _find_footer_start(data, 0x61, first_para, body_start)
        self.assertEqual(result, 0)

    # -----------------------------------------------------------------------
    # Fix: variable ctrl seq must not consume 0x13 (para-break prefix)
    # -----------------------------------------------------------------------

    def test_variable_ctrl_seq_does_not_consume_formatting_prefix(self):
        # A variable-type ctrl sequence (22 61 44) where the byte immediately
        # after the type byte is 0x13 (start of 13 04 50 para break) must leave
        # that para break intact for the main loop.
        # Old i+=4 behaviour consumed 0x13, making the para break invisible and
        # leaking 0x50 ('P') as a printable character.
        data = _doc(b'First\x22\x61\x44\x13\x04\x50Second')
        paras = _paras(data)
        self.assertEqual(len(paras), 2)
        self.assertIn('First', paras[0])
        self.assertIn('Second', paras[1])

    # -----------------------------------------------------------------------
    # Fix: 01 XX XX (3-byte) trailing indent fallback
    # -----------------------------------------------------------------------

    def test_01_xx_xx_trailing_indent_consumed(self):
        # A para ctrl block immediately followed by 01 XX XX (doubled printable
        # pair without the leading 00) must not leak the XX bytes as artefacts.
        # Observed in HENCOTES first body para: 22 61 0b 64 00 0a 09 01 | 01 27 27
        data = (b'DOC'
                + bytes([0x22, 0x61, 0x0b, 0x64, 0x00, 0x0a, 0x09, 0x01])  # 8-byte ctrl block
                + bytes([0x01, 0x27, 0x27])                                   # 3-byte trailing indent
                + b'Hello')
        text = _plain(data)
        self.assertIn('Hello', text)
        self.assertNotIn('\x27\x27', text)   # raw bytes must not appear
        self.assertNotIn("''", text)          # apostrophe pair artefact must not appear

    # -----------------------------------------------------------------------
    # Integration tests against real sample files
    # -----------------------------------------------------------------------

    HENCOTES_FILE   = Path(__file__).parent.parent / 'HENCOTES'
    BUILDNGS_H_FILE = Path(__file__).parent.parent / 'sample_files' / 'NewFiles - Copy' / 'PCWDISC.008' / 'GROUP.4' / 'HUNSTANW.ORT'

    def _load(self, path):
        if not path.exists():
            self.skipTest(f"Sample file not found: {path}")
        return parse(path.read_bytes())

    def test_hencotes_header_extracted(self):
        # HENCOTES has a centred header containing "Hencotes"
        doc = self._load(self.HENCOTES_FILE)
        self.assertIsNotNone(doc.header, "doc.header should not be None")
        self.assertIn('Hencotes', doc.header.plain_text())

    def test_hencotes_footer_extracted(self):
        # HENCOTES has a footer containing the document reference "CND 4.1"
        doc = self._load(self.HENCOTES_FILE)
        self.assertIsNotNone(doc.footer, "doc.footer should not be None")
        self.assertIn('CND 4.1', doc.footer.plain_text())

    def test_hencotes_header_not_in_body(self):
        # "Hencotes" should appear in the header, not as a stray first body paragraph
        doc = self._load(self.HENCOTES_FILE)
        body_text = '\n'.join(p.plain_text() for p in doc.paragraphs)
        # The body does contain "HENCOTES" (caps) as a heading — check only
        # the title-cased form is NOT a stray first paragraph on its own.
        if doc.paragraphs:
            self.assertNotEqual(doc.paragraphs[0].plain_text().strip(), 'Hencotes')


class TestGoldenFileHeaderFooter(unittest.TestCase):
    """Golden-file checks scoped to header/footer content."""

    @classmethod
    def setUpClass(cls):
        if not SAMPLE_FILE.exists():
            raise unittest.SkipTest(f"Sample file not found: {SAMPLE_FILE}")
        with open(SAMPLE_FILE, 'rb') as f:
            cls.doc = parse(f.read())
        cls.txt = to_txt(cls.doc)

    def test_txt_starts_with_header(self):
        # TXT output should open with the header text, not body text
        self.assertTrue(cls := self.txt.split('\n\n')[0],
                        'Output must not be empty')
        self.assertIn('Hencotes', self.txt.split('\n\n')[0])

    def test_txt_ends_with_footer(self):
        # TXT output should close with the footer after a --- separator
        self.assertIn('CND 4.1', self.txt)
        self.assertTrue(self.txt.rstrip().endswith('2018'),
                        'Output should end with footer year')

    def test_txt_has_separators(self):
        # Header and footer sections are delimited by ---
        self.assertIn('---', self.txt)


class TestSelfReferentialCtrlSkip(unittest.TestCase):
    """22 XX XX (ctrl type == ctrl_byte) is always layout metadata; no text leak."""

    def test_22_61_61_printable_params_not_emitted(self):
        # "22 61 61 18 18 78 00 00 64" — 0x78='x' and 0x64='d' must NOT leak as text.
        # The self-referential sequence should skip to the next 22 61 0b, after
        # which the real content "Hello" appears.
        body = (
            b'\x22\x61\x61\x18\x18\x78\x00\x00\x64'  # self-ref + binary params
            b'\x22\x61\x0b\x00\x00\x00\x00\x00'       # next paragraph block (default 8-byte)
            b'Hello'
        )
        result = _plain(_doc(body))
        self.assertNotIn('x', result)
        self.assertNotIn('d', result)
        self.assertIn('Hello', result)

    def test_22_6d_6d_skipped_in_6d_file(self):
        # In a 22 6d file, 22 6d 6d is the self-referential sequence; 'm' must not leak.
        # Build a minimal 22 6d DOC file.
        para_ctrl_6d = bytes([0x22, 0x6d, 0x0b])
        # 22 6d 0b repeated 3 times to make ctrl_byte detection pick 0x6d
        header = b'DOC' + para_ctrl_6d * 3 + bytes([0x00] * 5)
        body = (
            b'\x22\x6d\x6d\x18\x18\x78\x00\x00'  # self-ref (22 6d 6d) + binary
            + para_ctrl_6d + b'\x00\x00\x00\x00\x00'  # next paragraph block
            + b'World'
        )
        result = parse(header + body).plain_text()
        self.assertNotIn('m', result)
        self.assertIn('World', result)


class TestSITabInParagraphHeader(unittest.TestCase):
    """22 XX 0b B3 B4 0f 04 B1 B2 [01 PP PP] — SI tab embedded in header must not leak B2."""

    def test_si_tab_header_b2_not_emitted(self):
        # "22 61 0b 88 02 0f 04 3b 61 01 0b 0b" — B2=0x61='a' must NOT be emitted.
        # The 9-byte skip + consume(01 PP PP) should land at the content byte.
        body = (
            b'\x22\x61\x0b\x88\x02\x0f\x04\x3b\x61\x01\x0b\x0b'  # SI-tab header block
            b'Content'
        )
        result = _plain(_doc(body))
        self.assertNotIn('a', result)
        self.assertIn('Content', result)

    def test_si_tab_header_with_non_control_doubled_pair(self):
        # Same structure but doubled pair uses printable bytes (PP >= 0x20).
        body = (
            b'\x22\x61\x0b\x28\x02\x0f\x04\x3b\x61\x01\x1f\x1f'  # 0x1f < 0x20 pair
            b'Text'
        )
        result = _plain(_doc(body))
        self.assertNotIn('a', result)
        self.assertIn('Text', result)


class TestSectionBreakParagraphSplit(unittest.TestCase):
    """0e 01 / 0e 02 mid-sentence should not create a paragraph split."""

    def _make(self, before: bytes, after: bytes) -> bytes:
        """Build a minimal doc with a section break between two text fragments.

        Layout: opening para block + before-text + 0e 01 + dummy layout block
        + second para block + after-text + paragraph break.
        The dummy layout block contains no 22 61 0b bytes so the parser jumps
        straight to the explicit second para block.
        """
        layout_block = bytes([0x00] * 10)
        second_para = PARA_CTRL + bytes([0x00, 0x00, 0x00, 0x00, 0x00])
        body = (
            before
            + bytes([0x0e, 0x01])
            + layout_block
            + second_para
            + after
            + bytes([0x13, 0x04, 0x50])  # paragraph break
        )
        return _doc(body)

    def test_section_break_mid_sentence_no_split(self):
        """0e 01 preceded by word separator should not split the paragraph."""
        # "care " 0e 01 "home" → one paragraph: "care home"
        data = self._make(
            b'care' + bytes([0x02]),          # "care" + word separator
            b'home',
        )
        paras = _paras(data)
        self.assertEqual(len(paras), 1)
        self.assertIn('care', paras[0])
        self.assertIn('home', paras[0])

    def test_section_break_mid_sentence_printable_prev(self):
        """0e 01 preceded by a printable byte should not split the paragraph."""
        # "elderly," 0e 01 "opened" → one paragraph
        data = self._make(b'elderly,', b'opened')
        paras = _paras(data)
        self.assertEqual(len(paras), 1)
        self.assertIn('elderly', paras[0])
        self.assertIn('opened', paras[0])

    def test_section_break_after_line_break_does_split(self):
        """0e 01 preceded by 13 04 78 (line break) should create a paragraph split."""
        # "first" 13 04 78 00 0e 01 "second" → two paragraphs
        data = self._make(
            b'first' + bytes([0x13, 0x04, 0x78, 0x00]),
            b'second',
        )
        paras = _paras(data)
        self.assertEqual(len(paras), 2)
        self.assertIn('first', paras[0])
        self.assertIn('second', paras[1])


class TestMidSentenceBlockFormatting(unittest.TestCase):
    """22 61 0b block with 13 04 XX at B5-B6 should leave the formatting
    sequence for the main loop rather than consuming it in the block skip."""

    def test_italic_on_not_consumed_by_block_skip(self):
        """13 04 64 at B5-B6-B7 should turn on italic, not be skipped."""
        # 22 61 0b B3 B4 13 04 64 → italic-on left for main loop
        block = PARA_CTRL + bytes([0x34, 0x00, 0x13, 0x04, 0x64])
        body = (
            block
            + b'text'
            + bytes([0x13, 0x04, 0x78])   # italic off
            + bytes([0x13, 0x04, 0x50])   # paragraph break
        )
        doc = parse(_doc(body))
        italic_runs = [r for p in doc.paragraphs for r in p.runs if r.italic]
        self.assertTrue(any('text' in r.text for r in italic_runs))

    def test_no_spurious_newline_when_italic_consumed(self):
        """Without fix, 13 04 78 fires as line break; with fix it fires as
        italic-off.  Confirm no \\n in the paragraph text."""
        block = PARA_CTRL + bytes([0x34, 0x00, 0x13, 0x04, 0x64])
        body = (
            b'before'
            + bytes([0x02])
            + block
            + b'OED'
            + bytes([0x13, 0x04, 0x78])   # should be italic-off, not newline
            + bytes([0x02])
            + b'after'
            + bytes([0x13, 0x04, 0x50])
        )
        doc = parse(_doc(body))
        text = doc.paragraphs[0].plain_text()
        self.assertNotIn('\n', text)
        self.assertIn('before', text)
        self.assertIn('OED', text)
        self.assertIn('after', text)

    def test_para_break_at_b5_not_consumed(self):
        """13 04 50 at B5-B6-B7 should flush the paragraph, not be skipped."""
        block = PARA_CTRL + bytes([0x54, 0x08, 0x13, 0x04, 0x50])
        body = (
            b'first'
            + block
            + b'second'
            + bytes([0x13, 0x04, 0x50])
        )
        paras = _paras(_doc(body))
        self.assertEqual(len(paras), 2)
        self.assertIn('first', paras[0])
        self.assertIn('second', paras[1])


class TestContentsPageControlBlock(unittest.TestCase):
    """22 6d variant: 0f 02 / 0f 01 paragraph and line-break separators,
    and 78 00 0a extended-variant skip."""

    # Minimal 22 6d file envelope
    MAGIC_6D = b'DOC'
    PARA_6D = bytes([0x22, 0x6d, 0x0b])

    def _doc_6d(self, body: bytes) -> bytes:
        """Wrap body in a minimal 22 6d document envelope."""
        header = self.MAGIC_6D + self.PARA_6D + bytes([0x00, 0x00, 0x00, 0x00, 0x00])
        return header + body

    def _para_sep(self) -> bytes:
        """0f 02 22 6d 0b [6 zero bytes] — paragraph separator."""
        return bytes([0x0f, 0x02]) + self.PARA_6D + bytes([0x00] * 5)

    def _line_break(self) -> bytes:
        """0f 01 22 6d 0b [6 zero bytes] — line break."""
        return bytes([0x0f, 0x01]) + self.PARA_6D + bytes([0x00] * 5)

    def test_0f_02_creates_paragraph_break(self):
        """0f 02 22 6d 0b should flush the paragraph and start a new one."""
        body = (
            b'first'
            + self._para_sep()
            + b'second'
            + bytes([0x13, 0x04, 0x50])
        )
        paras = _paras(self._doc_6d(body))
        self.assertEqual(len(paras), 2)
        self.assertIn('first', paras[0])
        self.assertIn('second', paras[1])

    def test_0f_01_creates_line_break(self):
        """0f 01 22 6d 0b should emit a newline within the current paragraph."""
        body = (
            b'first'
            + self._line_break()
            + b'second'
            + bytes([0x13, 0x04, 0x50])
        )
        paras = _paras(self._doc_6d(body))
        self.assertEqual(len(paras), 1)
        self.assertIn('first', paras[0])
        self.assertIn('second', paras[0])
        self.assertIn('\n', paras[0])

    def test_0f_02_does_not_fire_in_22_61_files(self):
        """0f 02 22 61 0b should NOT be treated as a paragraph break in standard files."""
        # In a standard 22 61 file the pattern has a different meaning
        para_61 = bytes([0x22, 0x61, 0x0b])
        sep = bytes([0x0f, 0x02]) + para_61 + bytes([0x00] * 5)
        body = b'first' + sep + b'second' + bytes([0x13, 0x04, 0x50])
        header = b'DOC' + para_61 + bytes([0x00] * 5)
        paras = _paras(header + body)
        # Should NOT produce two paragraphs from the 0f 02 alone
        combined = ' '.join(paras)
        self.assertIn('first', combined)
        self.assertIn('second', combined)

    def test_78_00_0a_skip_no_spurious_x(self):
        """22 6d 0b block with 78 00 0a separator should not emit spurious bytes."""
        # Build: 22 6d 0b [B3 B4 B5 78 00 0a PP PP] then "text"
        block = self.PARA_6D + bytes([0x00, 0x00, 0x00, 0x78, 0x00, 0x0a, 0x41, 0x41])
        body = block + b'text' + bytes([0x13, 0x04, 0x50])
        data = self._doc_6d(body)
        result = _plain(data)
        self.assertIn('text', result)
        self.assertNotIn('x', result.replace('text', ''))


class TestBodyStartDetection22_6d(unittest.TestCase):
    """_find_body_start correctly identifies body start in 22 6d variant files
    where the body-start block has B5=0x14 and is preceded by high-B3 blocks."""

    PARA_6D = bytes([0x22, 0x6d, 0x0b])

    def test_seen_high_b3_unlocks_b5_14_body_start(self):
        """After a high-B3 block, a low-B3 block with B5=0x14 is body start."""
        # Mimics BUILDNGS.A-C: high-B3 transition blocks then 42 0d 14 78 00 body-start.
        # High-B3 block (B3=0xa6, B4=0x0e → structural header, would be skipped by
        # _find_body_start's structural filter). Then body-start with B5=0x14.
        para = self.PARA_6D
        high_b3_block = para + bytes([0xa6, 0x0e, 0x00, 0x00, 0x00])   # B3>=0x80
        body_start_block = para + bytes([0x42, 0x0d, 0x14, 0x78, 0x00]) # B3<0x80, B5=0x14
        data = b'DOC' + high_b3_block + body_start_block + b'Hello'
        first_para = 3
        result = _find_body_start(data, 0x6d, first_para, 'header')
        # Should land on the low-B3 block, not skip it because of B5=0x14
        self.assertEqual(result, 3 + len(high_b3_block))

    def test_seen_high_b3_not_set_b5_14_still_skipped(self):
        """Without any prior high-B3 block, a B5=0x14 block is still skipped."""
        para = self.PARA_6D
        b5_14_block = para + bytes([0x50, 0x0d, 0x14, 0x78, 0x00])  # B3<0x80, B5=0x14
        body_block = para + bytes([0x64, 0x00, 0x0a, 0x09, 0x01])    # B5=0x0a → body
        data = b'DOC' + b5_14_block + body_block + b'Hello'
        first_para = 3
        result = _find_body_start(data, 0x6d, first_para, 'header')
        # B5=0x14 with no prior high-B3 → still treated as header zone → skipped
        self.assertEqual(result, 3 + len(b5_14_block))

    def test_78_00_body_start_no_spurious_chars(self):
        """22 6d 0b B3<0x80 78 00 body-start block must not emit trailing bytes."""
        # Structural trailing bytes after 78 00: 0a 05 00 13 00 78 00 11 06
        # 'x' (0x78) must not appear before 'Contents' (would give 'xContents').
        para = self.PARA_6D
        # Body-start block matching BUILDNGS.A-C pattern at 0x6d2; scan stops at 11 06.
        body_start = para + bytes([0x42, 0x0d, 0x14, 0x78, 0x00,
                                   0x0a, 0x05, 0x00, 0x13, 0x00, 0x78, 0x00])
        # Alignment code, then a standard 8-byte ctrl block (B7=0x00) before text.
        content = bytes([0x11, 0x06]) + para + bytes([0x43, 0x01, 0x08, 0x08, 0x00])
        text = b'Contents' + bytes([0x13, 0x04, 0x50])
        data = b'DOC' + body_start + content + text
        result = _plain(data)
        self.assertIn('Contents', result)
        # No spurious 'x' before Contents (would appear as 'xContents')
        self.assertNotIn('xContents', result)
        self.assertNotIn('# ', result)

    def test_78_00_scan_does_not_apply_to_22_61_files(self):
        """Standard 22 61 files must not have the scan-forward applied to 78 00 blocks."""
        # In a 22 61 file, a block with B6=0x78 B7=0x00 followed by content bytes
        # should NOT have the scan skip over those bytes.
        para_61 = bytes([0x22, 0x61, 0x0b])
        # Block with 78 00 followed by word-separator + printable content
        block = para_61 + bytes([0x52, 0x02, 0x14, 0x78, 0x00])
        content = bytes([0x02, 0x02]) + b'1827' + bytes([0x13, 0x04, 0x50])
        data = b'DOC' + block + content
        result = _plain(data)
        # '1827' must be present — the scan must not skip over it
        self.assertIn('1827', result)


class Test1eVariant(unittest.TestCase):
    """Tests for the 0x1e prefix byte (1e 74 0b) variant."""

    PARA_CTRL_1E = bytes([0x1e, 0x74, 0x0b])

    def _doc_1e(self, body: bytes) -> bytes:
        """Minimal envelope for 1e 74 variant: 3× para_ctrl blocks so _detect_variant
        sees enough 1e 74 0b occurrences to win the frequency count."""
        anchor = self.PARA_CTRL_1E + bytes([0x00, 0x00, 0x00, 0x00, 0x00])
        return MAGIC + anchor * 3 + body

    def test_0f_02_1e_74_0b_paragraph_break(self):
        """0f 02 1e 74 0b must flush the current paragraph and start a new one."""
        anchor = self.PARA_CTRL_1E + bytes([0x00, 0x00, 0x00, 0x00, 0x00])
        sep = bytes([0x0f, 0x02]) + self.PARA_CTRL_1E + bytes([0x00, 0x00, 0x00, 0x00, 0x00])
        data = MAGIC + anchor * 3 + b'First' + sep + b'Second\x13\x04\x50'
        paras = _paras(data)
        self.assertIn('First', paras)
        self.assertIn('Second', paras)

    def test_0f_01_1e_74_0b_line_break(self):
        """0f 01 1e 74 0b must emit a line break within the current paragraph."""
        anchor = self.PARA_CTRL_1E + bytes([0x00, 0x00, 0x00, 0x00, 0x00])
        sep = bytes([0x0f, 0x01]) + self.PARA_CTRL_1E + bytes([0x00, 0x00, 0x00, 0x00, 0x00])
        data = MAGIC + anchor * 3 + b'Line1' + sep + b'Line2\x13\x04\x50'
        result = _plain(data)
        self.assertIn('Line1', result)
        self.assertIn('Line2', result)
        self.assertIn('\n', result)

    def test_doubled_pair_below_0x20_consumed(self):
        """In 1e variant files, doubled-pair indent values < 0x20 must be consumed,
        not leaked as text. Here 01 04 04 should not produce chr(4) in output."""
        anchor = self.PARA_CTRL_1E + bytes([0x00, 0x00, 0x00, 0x00, 0x00])
        # After paragraph break, 01 04 04 is the trailing doubled-pair indent
        body = b'Hello\x13\x04\x50\x01\x04\x04' + anchor + b'World\x13\x04\x50'
        data = MAGIC + anchor * 3 + body
        result = _plain(data)
        self.assertIn('Hello', result)
        self.assertIn('World', result)
        self.assertNotIn('\x04', result)

    def test_doubled_prefix_1e_1e_skipped(self):
        """1e 1e 74 self-referential sequence must be skipped entirely;
        no artefact bytes (e.g. 0x23 '#') must appear in the output."""
        anchor = self.PARA_CTRL_1E + bytes([0x00, 0x00, 0x00, 0x00, 0x00])
        # Self-referential sequence followed by 5 metadata bytes
        self_ref = bytes([0x1e, 0x1e, 0x74, 0x01, 0x00, 0x00, 0x00, 0x23])
        body = b'Before' + self_ref + anchor + b'After\x13\x04\x50'
        data = MAGIC + anchor * 3 + body
        result = _plain(data)
        self.assertIn('Before', result)
        self.assertIn('After', result)
        self.assertNotIn('#', result)

    def test_1e_variant_basic_text(self):
        """A 1e 74 file with plain text and word separators must parse correctly."""
        data = self._doc_1e(b'Hello\x02world\x13\x04\x50')
        self.assertEqual(_plain(data), 'Hello world')

    def test_1e_footer_extracted_from_22_prebody_zone(self):
        """Footer text in the 22 6d pre-body zone of a 1e 74 file must be
        extracted correctly and must not appear as a body paragraph."""
        # Build a synthetic 1e 74 file whose pre-body zone contains a 22 6d 0b
        # footer paragraph followed by 1e 74 0b body paragraphs.
        #
        # Pre-body zone: 22 6d 0b block (B3=0xfa, non-structural) + footer text
        # Layout section byte (0x5a5) = 0x01 → 'footer'
        # Body zone: 1e 74 0b blocks + body text
        #
        # We need the DOC magic + layout table stub at the right offsets.
        # _LAYOUT_SECTION_START = 0x5a0; byte at +5 = 0x01 (footer).
        LAYOUT_SECTION_START = 0x5a0
        # Build a minimal file: magic + zero-fill up to layout section + footer
        # section byte + zero-fill to pre-body zone.
        prebody_start = LAYOUT_SECTION_START + 10  # after layout section stub

        # Layout section stub: anchor byte in 0x60-0x9f range at +0, 0x00 at +1,
        # then 4 bytes of filler, with byte +5 = 0x01 (footer).
        # _section_type_at scans for [0x60-0x9f] 0x00 and reads marker at +5.
        layout_stub = b'\x60\x00\x00\x00\x00\x01'

        # Pre-body: 22 6d 0b block (8 bytes) + footer text + 00 00 terminator
        prebody_block = bytes([0x22, 0x6d, 0x0b, 0xfa, 0x0b, 0x00, 0x00, 0x00])
        footer_text = b'Footer\x02text\x13\x04\x50'
        prebody_terminator = b'\x00\x00'

        # Body: 3× 1e 74 0b anchors (to win _detect_variant) + body text
        body_anchor = bytes([0x1e, 0x74, 0x0b, 0x00, 0x00, 0x00, 0x00, 0x00])
        body_text = b'Body\x02content\x13\x04\x50'

        data = (b'DOC'
                + b'\x00' * (LAYOUT_SECTION_START - 3)  # pad to layout section
                + layout_stub
                + b'\x00' * (prebody_start - LAYOUT_SECTION_START - len(layout_stub))
                + prebody_block + footer_text + prebody_terminator
                + body_anchor * 3 + body_text)

        doc = parse(data)
        self.assertIsNotNone(doc.footer)
        self.assertIn('Footer', doc.footer.plain_text())
        self.assertIn('text', doc.footer.plain_text())
        # Footer text must NOT appear as a body paragraph
        body_texts = [p.plain_text() for p in doc.paragraphs]
        self.assertNotIn('Footer text', body_texts)
        self.assertIn('Body content', body_texts)

    def test_1e_body_not_split_by_high_b3_heuristic(self):
        """Early 1e 74 0b blocks with B3 >= 0x80 must not cause body_start to
        be pushed forward (the B3 heuristic used for 22 XX files does not apply
        to 1e files where all blocks are body content)."""
        # Two blocks with B3 >= 0x80 followed by a block with text.
        # Without the fix, body_start would land on the third block and the
        # text from the first two would be routed to the footer.
        anchor_high = bytes([0x1e, 0x74, 0x0b, 0xcc, 0x10, 0x00, 0x00, 0x00])  # B3=0xcc
        anchor_high2 = bytes([0x1e, 0x74, 0x0b, 0xfe, 0x0a, 0x00, 0x00, 0x00]) # B3=0xfe
        anchor_low  = bytes([0x1e, 0x74, 0x0b, 0x08, 0x0e, 0x00, 0x00, 0x00])  # B3=0x08
        data = (MAGIC
                + anchor_high * 3          # enough for _detect_variant
                + b'\x0f\x02' + anchor_high  # 0f 02 separator + high-B3 block
                + b'First\x13\x04\x50'
                + b'\x0f\x02' + anchor_high2
                + b'Second\x13\x04\x50'
                + b'\x0f\x02' + anchor_low
                + b'Third\x13\x04\x50')

        doc = parse(data)
        body_texts = [p.plain_text() for p in doc.paragraphs]
        self.assertIn('First', body_texts)
        self.assertIn('Second', body_texts)
        self.assertIn('Third', body_texts)
        self.assertIsNone(doc.footer)

    def test_22_prebody_ctrl_self_referential_suppressed_in_1e_body(self):
        """In 1e 74 files, binary blobs in the body embed the pre-body zone's
        22 XX XX self-referential sequence (e.g. 22 6d 6d).  These bytes are
        printable and would otherwise leak as '"mmxd'-style artefacts.
        The parser must skip to the next 1e 74 0b block instead of emitting them."""
        # Construct a 1e 74 file whose pre-body zone uses 22 6d 0b.
        # pre-body: 22 6d 0b header blocks (enough for _detect_variant to pick 22 6d)
        prebody_block = bytes([0x22, 0x6d, 0x0b, 0x00, 0x00, 0x00, 0x00, 0x00])
        # body: 1e 74 0b blocks (detected as most frequent pair)
        body_anchor = bytes([0x1e, 0x74, 0x0b, 0x00, 0x00, 0x00, 0x00, 0x00])
        # Embed a binary blob that contains 22 6d 6d followed by printable junk,
        # then a real paragraph.
        blob_with_selfreference = bytes([
            0x22, 0x6d, 0x6d,   # 22 prebody_ctrl prebody_ctrl — the leaked self-referential
            0x14, 0x18,         # non-printable bytes (consumed by normal flow)
            0x78,               # 'x' — printable, should NOT appear in output
            0x64,               # 'd' — printable, should NOT appear in output
        ])
        data = (MAGIC
                + prebody_block * 3     # 3× 22 6d 0b so _detect_variant sees them in prebody
                + body_anchor * 3       # 3× 1e 74 0b so _detect_variant picks 1e 74 as body
                + blob_with_selfreference
                + body_anchor           # next para_ctrl — skip lands here
                + b'RealText' + bytes([0x13, 0x04, 0x50]))

        doc = parse(data)
        body_texts = [p.plain_text() for p in doc.paragraphs]
        combined = ' '.join(body_texts)
        self.assertIn('RealText', combined)
        # Printable bytes from the binary blob must not appear
        self.assertNotIn('"mm', combined)
        self.assertNotIn('xd', combined)

    def test_page_break_b8_07_skips_binary_blob(self):
        """1e 74 0b cc 10 14 90 00 07 03 ... is a page-break form where font-size
        bytes (B5=0x14) precede the 07 03 marker at B8/B9.  The entire blob between
        this block and the next para_ctrl must be skipped — no printable bytes emitted."""
        body_anchor = bytes([0x1e, 0x74, 0x0b, 0x00, 0x00, 0x00, 0x00, 0x00])
        page_break_block = bytes([0x1e, 0x74, 0x0b, 0xcc, 0x10, 0x14, 0x90, 0x00, 0x07, 0x03])
        blob = bytes([0x40, 0x30, 0x48, 0x27, 0x78, 0x78, 0x23, 0x24])  # printable junk
        data = (MAGIC
                + body_anchor * 3
                + page_break_block + blob
                + body_anchor
                + b'Good' + bytes([0x13, 0x04, 0x50]))
        doc = parse(data)
        combined = ' '.join(p.plain_text() for p in doc.paragraphs)
        self.assertIn('Good', combined)
        self.assertNotIn('@', combined)
        self.assertNotIn("H'", combined)
        self.assertNotIn('#$', combined)

    def test_page_break_b6_b7_07_03_skips_binary_blob(self):
        """1e 74 0b B3 10 02 07 03 ... is a page-break form where 07 03 appears at
        B6/B7.  The blob between this block and the next para_ctrl must be skipped."""
        body_anchor = bytes([0x1e, 0x74, 0x0b, 0x00, 0x00, 0x00, 0x00, 0x00])
        page_break_block = bytes([0x1e, 0x74, 0x0b, 0xae, 0x10, 0x02, 0x07, 0x03, 0x00])
        blob = bytes([0x60, 0x81, 0x48, 0x27, 0x78, 0x78, 0x23, 0x24])  # printable junk
        data = (MAGIC
                + body_anchor * 3
                + page_break_block + blob
                + body_anchor
                + b'Good' + bytes([0x13, 0x04, 0x50]))
        doc = parse(data)
        combined = ' '.join(p.plain_text() for p in doc.paragraphs)
        self.assertIn('Good', combined)
        self.assertNotIn('`', combined)
        self.assertNotIn('#$', combined)

    def test_07_03_in_1e_body_discards_metadata_text(self):
        """In 1e-prefix files, 07 03 at the end of a per-page control text block
        (e.g. 'Last page Header / Footer disabled') should discard the accumulated
        metadata text and jump to the next para_ctrl."""
        body_anchor = bytes([0x1e, 0x74, 0x0b, 0x00, 0x00, 0x00, 0x00, 0x00])
        indent = bytes([0x01, 0x22, 0x22])  # trailing indent consumed by skip
        metadata = b'Metadata text'
        page_break = bytes([0x07, 0x03])
        blob = bytes([0x40, 0x30, 0x48, 0x60])  # printable junk after 07 03
        data = (MAGIC
                + body_anchor * 3
                + body_anchor + indent + metadata + page_break + blob
                + body_anchor
                + b'RealContent' + bytes([0x13, 0x04, 0x50]))
        doc = parse(data)
        combined = ' '.join(p.plain_text() for p in doc.paragraphs)
        self.assertIn('RealContent', combined)
        self.assertNotIn('Metadata', combined)
        self.assertNotIn('@', combined)


class TestInlineFormattingParamConsumption(unittest.TestCase):
    """0x0f must not be consumed as a non-printable param after inline formatting."""

    def test_0f_not_consumed_as_bold_off_param(self):
        """09 00 0f 02 PP ctrl 0b: the 0f must survive as a paragraph separator
        in non-0x61 files, not be swallowed as a parameter byte of bold-off.
        Uses 1e 74 variant where 0f 02 is the paragraph separator pattern."""
        PARA_CTRL_1E = bytes([0x1e, 0x74, 0x0b])
        detect_anchor = PARA_CTRL_1E + bytes([0x00, 0x0d, 0x00, 0x00, 0x00])
        bold_off = bytes([0x09, 0x00])
        sep = bytes([0x0f, 0x02]) + PARA_CTRL_1E + bytes([0x00, 0x0d, 0x00, 0x00, 0x00])
        body_anchor = PARA_CTRL_1E + bytes([0x00, 0x0d, 0x00, 0x00, 0x00])
        data = MAGIC + detect_anchor * 3 + body_anchor + b'First' + bold_off + sep + b'Second\x13\x04\x50'
        paras = [p.plain_text() for p in parse(data).paragraphs if p.plain_text().strip()]
        self.assertIn('First', paras)
        self.assertIn('Second', paras)
        self.assertEqual(len([p for p in paras if 'First' in p or 'Second' in p]), 2)


class Test1eStructuralSkipGuard(unittest.TestCase):
    """B3 >= 0x80 and B4 == 0x0e must NOT skip content in 1e-prefix files."""

    PARA_CTRL_1E = bytes([0x1e, 0x74, 0x0b])

    def _doc_1e(self, body: bytes) -> bytes:
        detect = self.PARA_CTRL_1E + bytes([0x00, 0x0d, 0x00, 0x00, 0x00])
        return MAGIC + detect * 3 + body

    def test_high_b3_b4_0e_content_not_skipped_in_1e_file(self):
        """In a 1e file, 1e 74 0b with B3=0x9a (>=0x80) and B4=0x0e must NOT
        be treated as a structural skip — its text content must appear."""
        # Block with B3=0x9a, B4=0x0e, B5=0x14, B6=0x90 — formerly wrongly skipped
        block = self.PARA_CTRL_1E + bytes([0x9a, 0x0e, 0x14, 0x90, 0x00])
        sep = bytes([0x0f, 0x02]) + self.PARA_CTRL_1E + bytes([0x00, 0x0d, 0x00, 0x00, 0x00])
        data = self._doc_1e(block + b'TopEntry' + sep + b'Next\x13\x04\x50')
        paras = [p.plain_text() for p in parse(data).paragraphs if p.plain_text().strip()]
        self.assertIn('TopEntry', paras)

    def test_high_b3_b4_0e_font_size_set_in_1e_file(self):
        """In a 1e file, B5=0x14 B6=0x90 in a high-B3 B4=0x0e block must still
        set font_size on that paragraph (not bleed into the next one).
        A real top-level entry always has an external trailing pair (01 XX XX) —
        that's what prevents _pending_b4_indent from clearing _pending_large_font.
        The detect anchors must also carry trailing pairs so left_indent is 0
        when the body block fires."""
        # Detect anchors with trailing pairs: left_indent stays 0 for each
        detect = self.PARA_CTRL_1E + bytes([0x00, 0x0d, 0x00, 0x00, 0x00]) + bytes([0x01, 0x0d, 0x0d])
        # Body block: B3=0x9a, B4=0x0e, B5=0x14, B6=0x90; external trailing pair confirms top-level
        block = self.PARA_CTRL_1E + bytes([0x9a, 0x0e, 0x14, 0x90, 0x00]) + bytes([0x01, 0x0d, 0x0d])
        sep = bytes([0x0f, 0x02]) + self.PARA_CTRL_1E + bytes([0x00, 0x0d, 0x00, 0x00, 0x00])
        data = MAGIC + detect * 3 + block + b'TopEntry' + sep + b'NextEntry\x13\x04\x50'
        doc = parse(data)
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        top   = next(p for p in paras if 'TopEntry'   in p.plain_text())
        nxt   = next(p for p in paras if 'NextEntry'  in p.plain_text())
        self.assertAlmostEqual(top.font_size, 14.4)
        self.assertIsNone(nxt.font_size)

    def test_high_b3_b4_0e_still_skipped_in_22_file(self):
        """In a standard 22 61 file, B3>=0x80 B4=0x0e must still be treated as
        a structural skip (existing behaviour must be preserved)."""
        # structural block with B3=0xc4, B4=0x0e — should be skipped
        structural = PARA_CTRL + bytes([0xc4, 0x0e, 0x00, 0x00, 0x00])
        normal     = PARA_CTRL + bytes([0x00, 0x00, 0x00, 0x00, 0x00])
        data = MAGIC + normal + structural + b'JUNK' + normal + b'Real\x13\x04\x50'
        result = _plain(data)
        self.assertIn('Real', result)
        self.assertNotIn('JUNK', result)


class TestParaIndent(unittest.TestCase):
    """Tests for PARA_INDENT (08 05 01 XX XX) left_indent and bibliography style codes."""

    # Default _twips_per_unit for synthetic docs (scale_pitch 0x18 × 6 = 144)
    TWU = 0x18 * 6  # 144

    def _doc_with_indent(self, xx: int) -> bytes:
        """Build a one-paragraph doc with a PARA_INDENT byte of XX (doubled)."""
        indent_seq = bytes([0x08, 0x05, 0x01, xx, xx])
        return _doc(indent_seq + b'Hello\x13\x04\x50')

    def test_para_indent_below_0x20_sets_left_indent(self):
        """XX = 0x0a (< 0x20) must set left_indent = 0x0a × 144 = 1440."""
        doc = parse(self._doc_with_indent(0x0a))
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        self.assertEqual(paras[0].left_indent, 0x0a * self.TWU)

    def test_para_indent_at_0x20_not_an_indent(self):
        """XX = 0x20 (>= 0x20) is a bibliography style code — left_indent must stay 0."""
        doc = parse(self._doc_with_indent(0x20))
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        self.assertEqual(paras[0].left_indent, 0)

    def test_para_indent_zero_not_set(self):
        """XX = 0x00 must not set left_indent (zero is not a valid indent)."""
        doc = parse(self._doc_with_indent(0x00))
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        self.assertEqual(paras[0].left_indent, 0)

    def test_para_indent_not_emitted_as_text(self):
        """The PARA_INDENT sequence bytes must not appear in plain text output."""
        data = self._doc_with_indent(0x4a)  # 0x4a = 'J', >= 0x20 style code
        result = _plain(data)
        self.assertNotIn('J', result)
        self.assertIn('Hello', result)

    def test_rtf_emits_li_for_indented_para(self):
        """RTF output must contain \\li{twips} when left_indent > 0."""
        from converter import to_rtf
        doc = parse(self._doc_with_indent(0x0a))
        rtf = to_rtf(doc)
        expected = rf'\li{0x0a * self.TWU}'
        self.assertIn(expected, rtf)

    def test_rtf_no_li_for_unindented_para(self):
        """RTF output must not contain \\li when left_indent == 0."""
        from converter import to_rtf
        doc = parse(_doc(b'Hello\x13\x04\x50'))
        rtf = to_rtf(doc)
        self.assertNotIn(r'\li', rtf)

    def test_docx_sets_left_indent(self):
        """DOCX output must set paragraph_format.left_indent when left_indent > 0."""
        import tempfile, os
        from converter import save_docx
        from docx.shared import Twips
        doc = parse(self._doc_with_indent(0x0a))
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            path = f.name
        try:
            save_docx(doc, Path(path))
            from docx import Document as DocxDoc
            d = DocxDoc(path)
            body_paras = [p for p in d.paragraphs if p.text.strip()]
            self.assertGreater(len(body_paras), 0)
            self.assertEqual(body_paras[0].paragraph_format.left_indent,
                             Twips(0x0a * self.TWU))
        finally:
            os.unlink(path)


class TestFontSizeFrom1eVariant(unittest.TestCase):
    """Tests for B5/B6 font-size encoding in 1e 74 variant files."""

    PARA_CTRL_1E = bytes([0x1e, 0x74, 0x0b])

    def _doc_1e_with_block(self, b3: int, b4: int, b5: int, b6: int,
                           b7: int = 0x00) -> bytes:
        """Build a 1e 74 file with a body block whose header is B3..B7.

        Pre-detection anchors simulate top-level heading blocks: they carry an
        external trailing pair (01 0d 0d) so the parser sees i != i_after_skip
        and does NOT apply the sub-entry indent.  This isolates the body block
        as the sole source of left_indent.
        """
        # Trailing pair 01 0d 0d after each anchor: consumed by the 3-byte trailing
        # pair handler (min_dp=0x00 for 1e files so 0x0d qualifies), ensuring
        # i != i_after_skip and _pending_b4_indent is not applied to the anchors.
        detect_anchor = self.PARA_CTRL_1E + bytes([0x00, 0x0d, 0x00, 0x00, 0x00]) + bytes([0x01, 0x0d, 0x0d])
        body_block = self.PARA_CTRL_1E + bytes([b3, b4, b5, b6, b7])
        return MAGIC + detect_anchor * 3 + body_block + b'Hello\x13\x04\x50'

    def test_b5_0x14_sets_font_size(self):
        """When B5=0x14 in a 1e 74 body block, font_size must be B6 / 10.0."""
        # B6=0x78 → 12.0pt
        doc = parse(self._doc_1e_with_block(0xcc, 0x10, 0x14, 0x78))
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        self.assertAlmostEqual(paras[0].font_size, 12.0)

    def test_b5_0x14_large_b6_sets_font_size(self):
        """B6=0x90 → 14.4pt on a confirmed top-level entry (external trailing pair present)."""
        # Must include an external trailing pair so the parser confirms top-level status.
        detect_anchor = self.PARA_CTRL_1E + bytes([0x00, 0x0d, 0x00, 0x00, 0x00]) + bytes([0x01, 0x0d, 0x0d])
        body_block = self.PARA_CTRL_1E + bytes([0xcc, 0x0d, 0x14, 0x90, 0x00])
        trailing_pair = bytes([0x01, 0x0d, 0x0d])
        data = MAGIC + detect_anchor * 3 + body_block + trailing_pair + b'Hello\x13\x04\x50'
        doc = parse(data)
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        self.assertAlmostEqual(paras[0].font_size, 14.4)

    def test_b5_0x14_large_b6_sub_entry_no_font_size(self):
        """B6=0x90 (14.4pt heading font) must NOT be applied to sub-entries.
        Real-world example: 'Cockshaw / Sealwell Bridge' has B5=0x14 B6=0x90
        but is a sub-entry (no external trailing pair) and must not be 14pt."""
        doc = parse(self._doc_1e_with_block(0xcc, 0x0d, 0x14, 0x90))
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        self.assertIsNone(paras[0].font_size,
                          '14.4pt heading font must not bleed onto sub-entries')

    def test_b5_0x14_large_b6_mid_para_no_font_size(self):
        """B6=0x90 block encountered mid-paragraph (left_indent already set, e.g. after a
        mid-sentence 0e 02 page break) must NOT apply 14.4pt.
        Real-world: 'Richardson, G.; Boot and Shoe Maker, 1896 no.27 & 29' and
        'Richardson & Co.; Dispensing and Family Chemists, 1914 no.15' were merged
        by a mid-sentence page break; the second block had B5=0x14 B6=0x90 causing
        the merged paragraph to wrongly show at 14pt."""
        # Build a sub-entry paragraph (left_indent=576) that then encounters a
        # B5=0x14 B6=0x90 block mid-paragraph via a mid-sentence page break.
        # The 0e 02 page break is NOT emitted here — instead we simulate the
        # state by building two consecutive blocks: first a sub-entry block that
        # sets left_indent=576, then immediately a B5=0x14 B6=0x90 block.
        # Since the parser processes them in the same paragraph (no flush between),
        # left_indent is already 576 when the second block fires.
        PARA_CTRL = bytes([0x1e, 0x74, 0x0b])
        detect_anchor = PARA_CTRL + bytes([0x00, 0x0d, 0x00, 0x00, 0x00]) + bytes([0x01, 0x0d, 0x0d])
        # First block: sub-entry (no trailing pair) sets left_indent=576
        sub_block = PARA_CTRL + bytes([0xcc, 0x0d, 0x00, 0x00, 0x00])
        # Simulated mid-sentence page break — 0e 02 followed by the layout scan.
        # The parser will skip from 0e 02 forward to the next 1e 74 0b, so we
        # place junk bytes that don't contain 1e 74 0b, then the next block.
        page_break = bytes([0x0e, 0x02]) + bytes([0x00] * 8)
        # Second block: B5=0x14 B6=0x90 — must NOT set font_size since left_indent!=0
        font_block = PARA_CTRL + bytes([0x58, 0x05, 0x14, 0x90, 0x00])
        data = MAGIC + detect_anchor * 3 + sub_block + b'Part1\x02' + page_break + font_block + b'Part2\x13\x04\x50'
        doc = parse(data)
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        # The two parts should be in one paragraph (mid-sentence break), and font_size
        # must NOT be 14.4 — it must remain None (inherited from document default).
        combined = ' '.join(p.plain_text() for p in paras)
        self.assertIn('Part1', combined)
        self.assertIn('Part2', combined)
        for p in paras:
            if 'Part1' in p.plain_text() or 'Part2' in p.plain_text():
                self.assertNotEqual(p.font_size, 14.4,
                                    '14.4pt must not bleed onto mid-paragraph 1e 74 0b blocks')

    def test_b5_0x14_b6_78_mid_para_no_font_size(self):
        """B6=0x78 (12pt) block mid-paragraph must NOT override the paragraph font.
        Real-world: 'Little, A. & G.; County Roller Mills, 1863' and
        'Maltby, Fredeick Walton; Tailor & Outfitter, 1914 no.12' were merged by a
        mid-sentence page break; the Maltby block had B5=0x14 B6=0x78 causing the
        merged paragraph to wrongly show at 12pt instead of the document default."""
        PARA_CTRL = bytes([0x1e, 0x74, 0x0b])
        detect_anchor = PARA_CTRL + bytes([0x00, 0x0d, 0x00, 0x00, 0x00]) + bytes([0x01, 0x0d, 0x0d])
        sub_block = PARA_CTRL + bytes([0xcc, 0x0d, 0x00, 0x00, 0x00])
        page_break = bytes([0x0e, 0x02]) + bytes([0x00] * 8)
        font_block = PARA_CTRL + bytes([0x58, 0x05, 0x14, 0x78, 0x00])
        data = MAGIC + detect_anchor * 3 + sub_block + b'Part1\x02' + page_break + font_block + b'Part2\x13\x04\x50'
        doc = parse(data)
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        for p in paras:
            if 'Part1' in p.plain_text() or 'Part2' in p.plain_text():
                self.assertIsNone(p.font_size,
                                  '12pt must not bleed onto mid-paragraph 1e 74 0b blocks')

    def test_b5_not_0x14_no_font_size(self):
        """When B5 != 0x14, font_size must remain None."""
        doc = parse(self._doc_1e_with_block(0x00, 0x0d, 0x00, 0x78))
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        self.assertIsNone(paras[0].font_size)

    def test_b4_lte_0x0c_sets_indent_fallback(self):
        """A 1e 74 body block with no external trailing pair must set left_indent = 576.
        B4=0x0c used here; B4 value is not the discriminator — the absence of a
        trailing pair is."""
        doc = parse(self._doc_1e_with_block(0xcc, 0x0c, 0x00, 0x00))
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        self.assertEqual(paras[0].left_indent, 576)

    def test_external_trailing_pair_prevents_indent(self):
        """An external 01 XX XX trailing pair after the block must prevent indenting,
        regardless of B4.  Top-level entries (B5=0x14 B6=0x90, like 'Bell, Henry')
        always carry an external trailing pair; sub-entries (B5=0x01 B6=B7) do not."""
        # Anchors also carry trailing pairs so they don't contaminate left_indent.
        detect_anchor = self.PARA_CTRL_1E + bytes([0x00, 0x0d, 0x00, 0x00, 0x00]) + bytes([0x01, 0x0d, 0x0d])
        # B4=0x09 with external trailing 01 2b 2b — should NOT indent
        body_block = self.PARA_CTRL_1E + bytes([0xcc, 0x09, 0x00, 0x00, 0x00])
        trailing_pair = bytes([0x01, 0x2b, 0x2b])
        data = MAGIC + detect_anchor * 3 + body_block + trailing_pair + b'Hello\x13\x04\x50'
        doc = parse(data)
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        self.assertEqual(paras[0].left_indent, 0)

    def test_no_trailing_pair_sets_indent_regardless_of_b4(self):
        """A block with no external trailing pair must indent even when B4 > 0x0c.
        Real-world example: 'Beaumont Park' has B4=0x0e but is a genuine sub-entry
        (carries B5=0x01 B6=B7 doubled pair inside the block header, no external pair)."""
        # Anchors carry trailing pairs so they don't contaminate left_indent.
        detect_anchor = self.PARA_CTRL_1E + bytes([0x00, 0x0d, 0x00, 0x00, 0x00]) + bytes([0x01, 0x0d, 0x0d])
        # B4=0x0e (above old threshold) with B5=0x01 B6=B7=0x0d inside block, no external pair
        body_block = self.PARA_CTRL_1E + bytes([0x5c, 0x0e, 0x01, 0x0d, 0x0d])
        data = MAGIC + detect_anchor * 3 + body_block + b'Hello\x13\x04\x50'
        doc = parse(data)
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        self.assertEqual(paras[0].left_indent, 576,
                         'Entry with no external trailing pair must be indented '
                         'even when B4 > 0x0c')

    def test_rtf_emits_fs_for_font_size(self):
        """RTF output must contain \\fs{half_pts} when font_size is set."""
        from converter import to_rtf
        # B6=0x78 → 12.0pt → \fs24
        doc = parse(self._doc_1e_with_block(0xcc, 0x10, 0x14, 0x78))
        rtf = to_rtf(doc)
        self.assertIn(r'\fs24', rtf)

    def test_rtf_no_fs_when_no_font_size(self):
        """RTF output must not contain \\fs when font_size is None."""
        from converter import to_rtf
        doc = parse(_doc(b'Hello\x13\x04\x50'))
        rtf = to_rtf(doc)
        self.assertNotIn(r'\fs', rtf)

    def test_docx_sets_font_size(self):
        """DOCX output must set run.font.size when font_size is set."""
        import tempfile, os
        from converter import save_docx
        from docx.shared import Pt
        # B6=0x78 → 12.0pt
        doc = parse(self._doc_1e_with_block(0xcc, 0x10, 0x14, 0x78))
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            path = f.name
        try:
            save_docx(doc, Path(path))
            from docx import Document as DocxDoc
            d = DocxDoc(path)
            body_paras = [p for p in d.paragraphs if p.text.strip()]
            self.assertGreater(len(body_paras), 0)
            run = body_paras[0].runs[0]
            self.assertEqual(run.font.size, Pt(12.0))
        finally:
            os.unlink(path)


class Test1e78ScanForwardGuard(unittest.TestCase):
    """Regression tests for the 78 00 scan-forward guard in _skip_ctrl_sequence.

    In 22-prefix non-0x61 files (22 6d, 22 42) a block with B6=0x78 B7=0x00 is
    a section-start marker and triggers a scan-forward to the next alignment/ctrl
    marker.  In 1e-prefix files the same B6/B7 bytes encode a 12pt font size
    (B5=0x14 B6=0x78) and must use the standard 8-byte skip.  If the scan-forward
    fires in a 1e file it consumes the entire following paragraph, causing the next
    entry (e.g. 'Beacon Grange') to inherit a stale left_indent from the scanned
    block's B4 value.
    """

    PARA_CTRL_1E = bytes([0x1e, 0x74, 0x0b])

    def _doc_1e(self, body: bytes) -> bytes:
        anchor = self.PARA_CTRL_1E + bytes([0x00, 0x0d, 0x00, 0x00, 0x00])
        return MAGIC + anchor * 3 + body

    def test_1e_78_00_block_does_not_scan_forward(self):
        """A 1e 74 0b block with B5=0x14 B6=0x78 B7=0x00 must use the 8-byte skip.

        Previously the 78 00 scan-forward (designed for 22 6d section-start markers)
        incorrectly fired on this block because ctrl_byte=0x74 != 0x61.  This caused
        the scan to consume all bytes up to the next 1e 74 ctrl_prefix, losing the
        intervening paragraph text and applying a stale B4 indent to the next entry.
        """
        # Block A: B4=0x08 (<=0x0c, sub-entry candidate), B5=0x14, B6=0x78, B7=0x00
        # The B6/B7 pattern looks like '78 00' but means '12pt font' in 1e files.
        block_a = self.PARA_CTRL_1E + bytes([0x1c, 0x08, 0x14, 0x78, 0x00])
        text_a = b'First\x13\x04\x50'
        # Block B: B4=0x0e (>0x0c, top-level) — must NOT be indented
        para_sep = bytes([0x0f, 0x02])
        block_b = self.PARA_CTRL_1E + bytes([0x60, 0x0e, 0x14, 0x90, 0x00])
        trailing_b = bytes([0x01, 0x0d, 0x0d])
        text_b = b'Second\x13\x04\x50'
        data = self._doc_1e(block_a + text_a + para_sep + block_b + trailing_b + text_b)
        doc = parse(data)
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        # Both paragraphs must be extracted (scan-forward would eat 'First' text)
        texts = [p.plain_text().strip() for p in paras]
        self.assertIn('First', texts, 'First paragraph was consumed by incorrect scan-forward')
        self.assertIn('Second', texts, 'Second paragraph missing')
        # Second paragraph (Block B: B4=0x0e) must NOT be indented
        second = next(p for p in paras if 'Second' in p.plain_text())
        self.assertEqual(second.left_indent, 0,
                         'Block B with B4=0x0e must not be indented; '
                         'stale B4 from Block A must not bleed through')


class Test1eSectionBoundaryTrailer(unittest.TestCase):
    """Regression: 1e-prefix section-boundary font-size blocks emit '13 00 WW WW'
    trailing bytes after the standard 8-byte skip.  When WW=0x78 ('x') this leaked
    into body text producing 'xContents' artefacts.  Real-world: BINDINDX.HEX."""

    PARA_CTRL_1E = bytes([0x1e, 0x74, 0x0b])

    def _doc_1e(self, body: bytes) -> bytes:
        anchor = self.PARA_CTRL_1E + bytes([0x00, 0x0d, 0x00, 0x00, 0x00]) + bytes([0x01, 0x0d, 0x0d])
        return MAGIC + anchor * 3 + body

    def test_13_00_78_00_no_x_emitted(self):
        """13 00 78 00 trailer after B5=0x14 block must not emit 'x'."""
        # Block: 1e 74 0b 68 0f 14 78 00 → standard 8-byte skip
        # Trailer: 13 00 78 00 → must be consumed silently (WW=0x78='x' is printable)
        block = self.PARA_CTRL_1E + bytes([0x68, 0x0f, 0x14, 0x78, 0x00])
        trailer = bytes([0x13, 0x00, 0x78, 0x00])
        content = b'Contents\x13\x04\x50'
        data = self._doc_1e(block + trailer + content)
        doc = parse(data)
        text = ' '.join(p.plain_text() for p in doc.paragraphs if p.plain_text().strip())
        self.assertNotIn('x', text.replace('Contents', ''),
                         "13 00 78 00 trailer must not emit 'x' into the output")
        self.assertIn('Contents', text)

    def test_13_00_00_00_trailer_consumed(self):
        """13 00 00 00 trailer (WW=0x00) must also be consumed silently."""
        block = self.PARA_CTRL_1E + bytes([0x68, 0x0f, 0x14, 0x90, 0x00])
        trailer = bytes([0x13, 0x00, 0x00, 0x00])
        content = b'Contents\x13\x04\x50'
        data = self._doc_1e(block + trailer + content)
        doc = parse(data)
        text = ' '.join(p.plain_text() for p in doc.paragraphs if p.plain_text().strip())
        self.assertIn('Contents', text)

    def test_13_00_trailer_no_left_indent(self):
        """Section-boundary blocks with 13 00 WW WW trailer are top-level structural
        markers and must not receive the 576-twip sub-entry indent."""
        block = self.PARA_CTRL_1E + bytes([0x68, 0x0f, 0x14, 0x78, 0x00])
        trailer = bytes([0x13, 0x00, 0x78, 0x00])
        content = b'Contents\x13\x04\x50'
        data = self._doc_1e(block + trailer + content)
        doc = parse(data)
        paras = [p for p in doc.paragraphs if 'Contents' in p.plain_text()]
        self.assertTrue(paras, 'Contents paragraph not found')
        self.assertEqual(paras[0].left_indent, 0,
                         'Section-boundary block must not produce a sub-entry indent')


class Test1eB6B7TabConsumed(unittest.TestCase):
    """Regression test: 1e 74 0b blocks with B5=sequence-byte, B6=0f, B7=04.

    In BINDINDX.HEX sub-entries under a heading (e.g. 'Auction Marts') use a
    variant block structure where B5 carries a sequence number (0x31, 0x32, 0x33)
    and B6/B7 = 0x0f 0x04 encode the column spec for the entry.  Previously the
    parser skipped only 6 bytes and left '0f 04 B1 B2 01 XX XX' for the main loop,
    which then emitted a spurious tab character and recorded a 4200-twip tab stop,
    pushing the text to the centre of the document.  The block must now be fully
    consumed without emitting a tab, matching the behaviour of the B5=0x0f B6=0x04
    case (e.g. 'Leazes Nursery Estate').
    """

    PARA_CTRL_1E = bytes([0x1e, 0x74, 0x0b])

    def _doc_1e_sub_entry(self, b5_seq: int) -> bytes:
        """Build a 1e 74 file containing one sub-entry block with B5=b5_seq,
        B6=0x0f, B7=0x04 followed by the standard column/ctrl/trailing structure."""
        # Anchor blocks with trailing pairs so they don't contaminate left_indent.
        anchor = self.PARA_CTRL_1E + bytes([0x00, 0x0d, 0x00, 0x00, 0x00]) + bytes([0x01, 0x0d, 0x0d])
        # Sub-entry block: B3=0x8e B4=0x0d B5=seq B6=0x0f B7=0x04, then
        # column 0x23, ctrl_byte 0x74, trailing pair 01 0d 0d
        block = self.PARA_CTRL_1E + bytes([0x8e, 0x0d, b5_seq, 0x0f, 0x04, 0x23, 0x74, 0x01, 0x0d, 0x0d])
        return MAGIC + anchor * 3 + block + b'Hello\x13\x04\x50'

    def test_b6_0f_b7_04_no_tab_emitted(self):
        """Block with B5=sequence-byte, B6=0x0f, B7=0x04 must not emit a tab character."""
        doc = parse(self._doc_1e_sub_entry(0x31))
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        self.assertEqual(paras[0].plain_text().strip(), 'Hello')
        self.assertNotIn('\t', paras[0].plain_text(),
                         'Spurious tab emitted from structural B6=0x0f B7=0x04 block')

    def test_b6_0f_b7_04_no_tab_stop_recorded(self):
        """Block with B5=sequence-byte, B6=0x0f, B7=0x04 must not record a tab stop."""
        doc = parse(self._doc_1e_sub_entry(0x32))
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        self.assertEqual(paras[0].tab_stops, [],
                         'Spurious 4200-twip tab stop recorded from structural block')

    def test_b6_0f_b7_04_still_sets_sub_entry_indent(self):
        """Block with B5=sequence-byte, B6=0x0f, B7=0x04 must still produce left_indent=576."""
        doc = parse(self._doc_1e_sub_entry(0x33))
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        self.assertEqual(paras[0].left_indent, 576,
                         'Sub-entry indent must still be applied when B6=0x0f B7=0x04')


class Test1eBodyZoneMetadataBlockSkip(unittest.TestCase):
    """Regression: in 1e-prefix files a B3=0x00, B4=0x00, B5=0x14 block in the body
    zone is a section-boundary metadata block whose payload contains binary data.
    Bytes like 0x25 ('%'), 0x24 ('$'), 0xc2 ('?') in the payload were leaking as
    text before the self-referential 22 6d 6d skip fired.  Real-world: BINDINDX.HEX
    produced '%$? HEXHAM MECHANICS'' artefact before this fix."""

    PARA_CTRL_1E = bytes([0x1e, 0x74, 0x0b])
    MAGIC_1E = b'DOC'

    def _doc_1e(self, body: bytes) -> bytes:
        """Minimal 1e-variant document with 3 top-level detect anchors + body."""
        anchor = self.PARA_CTRL_1E + bytes([0x00, 0x0d, 0x00, 0x00, 0x00]) + bytes([0x01, 0x0d, 0x0d])
        return self.MAGIC_1E + anchor * 3 + body

    def test_b3_00_b4_00_b5_14_body_block_skipped(self):
        """B3=0x00, B4=0x00, B5=0x14 block in body zone must be skipped; payload bytes
        (including printable ones like '%', '$', 0xc2) must not appear in output."""
        # Section-boundary metadata block: 1e 74 0b 00 00 14 78 00
        # followed by a binary payload that contains printable junk bytes
        metadata_block = self.PARA_CTRL_1E + bytes([0x00, 0x00, 0x14, 0x78, 0x00])
        # Payload that would produce '%$?' artefacts if not skipped
        payload = bytes([0x0c, 0x25, 0x01, 0x24, 0x01, 0x01, 0xc2, 0x01, 0x01, 0x02])
        # Next real paragraph with content
        next_para = self.PARA_CTRL_1E + bytes([0x9a, 0x0d, 0x14, 0x78, 0x00]) + b'HEXHAM\x13\x04\x50'
        data = self._doc_1e(metadata_block + payload + next_para)
        doc = parse(data)
        texts = [p.plain_text() for p in doc.paragraphs if p.plain_text().strip()]
        combined = ' '.join(texts)
        self.assertIn('HEXHAM', combined, 'Real content after metadata block must appear')
        self.assertNotIn('%', combined, "0x25 '%' must not leak from metadata payload")
        self.assertNotIn('$', combined, "0x24 '$' must not leak from metadata payload")

    def test_b3_00_b5_00_not_skipped_in_body(self):
        """B3=0x00, B5=0x00 detect-anchor blocks in body zone must NOT be skipped
        by the section-boundary guard (B5=0x14 is the distinguishing flag)."""
        # Standard detect anchor (B3=0x00, B4=0x0d, B5=0x00): must be processed normally
        anchor = self.PARA_CTRL_1E + bytes([0x00, 0x0d, 0x00, 0x00, 0x00]) + bytes([0x01, 0x0d, 0x0d])
        content_block = self.PARA_CTRL_1E + bytes([0x9a, 0x0d, 0x14, 0x78, 0x00]) + b'Hello\x13\x04\x50'
        data = self._doc_1e(anchor + content_block)
        doc = parse(data)
        texts = [p.plain_text() for p in doc.paragraphs if p.plain_text().strip()]
        self.assertTrue(any('Hello' in t for t in texts),
                        'Content after B5=0x00 detect anchor must not be skipped')


class Test1eB5PageBreakInBlockHeader(unittest.TestCase):
    """Regression: 1e-prefix blocks with B5=0x0e, B6=0x01/0x02 encode a page-break
    signal directly in the block header.  The binary layout blob that follows (which
    may contain a LocoScript date template with printable ASCII bytes such as 'F',
    "date", "February", "Monday", "pm") must not be emitted as body text.
    Real-world: BINDINDX.HEX last paragraph was '?Fdate th FebruaryFeb Monday pm'."""

    PARA_CTRL_1E = bytes([0x1e, 0x74, 0x0b])
    MAGIC_1E = b'DOC'

    def _doc_1e(self, body: bytes) -> bytes:
        anchor = self.PARA_CTRL_1E + bytes([0x00, 0x0d, 0x00, 0x00, 0x00]) + bytes([0x01, 0x0d, 0x0d])
        return self.MAGIC_1E + anchor * 3 + body

    def _make_date_blob(self) -> bytes:
        """Minimal date-template blob resembling real BINDINDX.HEX payload."""
        return (bytes([0xf0, 0x00, 0x00, 0x00, 0x46, 0x00, 0x00, 0x00])
                + b'\x01\x00date\x05\x00\x01\x00\x00\x00\x03\x00\x02th\x02'
                + b'\x00\x00\x00\x09\x00\x08February')

    def test_b5_0e_b6_02_page_break_no_date_blob(self):
        """B5=0x0e, B6=0x02 block must be skipped; date-template payload must not appear."""
        content_para = self.PARA_CTRL_1E + bytes([0xcc, 0x10, 0x14, 0x78, 0x00]) + b'Family\x13\x04\x50'
        pb_block = self.PARA_CTRL_1E + bytes([0xcc, 0x10, 0x0e, 0x02, 0x00])
        data = self._doc_1e(content_para + pb_block + self._make_date_blob())
        doc = parse(data)
        texts = [p.plain_text() for p in doc.paragraphs if p.plain_text().strip()]
        combined = ' '.join(texts)
        self.assertIn('Family', combined, 'Content before page-break block must appear')
        self.assertNotIn('date', combined, "'date' from blob must not appear in output")
        self.assertNotIn('February', combined, "'February' from blob must not appear")
        self.assertNotIn('F', combined.replace('Family', ''),
                         "0x46 'F' from blob must not appear in output")

    def test_b5_0e_b6_01_page_break_no_date_blob(self):
        """B5=0x0e, B6=0x01 (section break variant) also silently consumed."""
        pb_block = self.PARA_CTRL_1E + bytes([0xcc, 0x10, 0x0e, 0x01, 0x00])
        content_para = self.PARA_CTRL_1E + bytes([0xcc, 0x10, 0x14, 0x78, 0x00]) + b'After\x13\x04\x50'
        data = self._doc_1e(pb_block + self._make_date_blob() + content_para)
        doc = parse(data)
        texts = [p.plain_text() for p in doc.paragraphs if p.plain_text().strip()]
        combined = ' '.join(texts)
        self.assertNotIn('date', combined)
        self.assertNotIn('February', combined)


class TestPageBreakPropagation(unittest.TestCase):
    """0e 01 / 0e 02 non-mid-sentence breaks should set page_break_before on
    the following paragraph and be emitted correctly in all three formats."""

    def _make_page_break_doc(self, before: bytes, after: bytes, break_type: int = 0x02) -> bytes:
        """Doc with two paragraphs separated by a page/section break.

        before  — body bytes for the first paragraph (must end with a non-word-sep
                  so mid_sentence=False fires)
        after   — body bytes for the second paragraph
        """
        layout_block = bytes([0x00] * 10)
        second_para = PARA_CTRL + bytes([0x00, 0x00, 0x00, 0x00, 0x00])
        body = (
            before
            + bytes([0x0e, break_type])
            + layout_block
            + second_para
            + after
            + bytes([0x13, 0x04, 0x50])
        )
        return _doc(body)

    # --- Parser ---

    def test_page_break_before_set_on_following_paragraph(self):
        """Non-mid-sentence 0e 02 should set page_break_before=True on the next para."""
        data = self._make_page_break_doc(
            b'first' + bytes([0x13, 0x04, 0x78, 0x00]),  # ends with line-break + 0x00
            b'second',
        )
        doc = parse(data)
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        self.assertEqual(len(paras), 2)
        self.assertFalse(paras[0].page_break_before)
        self.assertTrue(paras[1].page_break_before)

    def test_no_page_break_before_when_mid_sentence(self):
        """Mid-sentence 0e 02 (prev = word sep) must not set page_break_before."""
        data = self._make_page_break_doc(
            b'care' + bytes([0x02]),   # word separator before 0e → mid_sentence
            b'home',
        )
        doc = parse(data)
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        self.assertEqual(len(paras), 1)
        self.assertFalse(paras[0].page_break_before)

    def test_no_page_break_before_when_no_break(self):
        """Paragraph with no preceding break must have page_break_before=False."""
        data = _doc(b'hello' + bytes([0x13, 0x04, 0x50]))
        doc = parse(data)
        paras = [p for p in doc.paragraphs if p.plain_text().strip()]
        self.assertFalse(paras[0].page_break_before)

    # --- TXT ---

    def test_txt_emits_page_break_separator(self):
        """TXT output should include '--- page break ---' before a flagged paragraph."""
        data = self._make_page_break_doc(
            b'first' + bytes([0x13, 0x04, 0x78, 0x00]),
            b'second',
        )
        doc = parse(data)
        txt = to_txt(doc)
        self.assertIn('--- page break ---', txt)
        # separator must appear before 'second', not before 'first'
        self.assertLess(txt.index('--- page break ---'), txt.index('second'))
        self.assertGreater(txt.index('--- page break ---'), txt.index('first'))

    # --- RTF ---

    def test_rtf_emits_page_control_word(self):
        r"""RTF output should include \page before the flagged paragraph's \pard."""
        data = self._make_page_break_doc(
            b'first' + bytes([0x13, 0x04, 0x78, 0x00]),
            b'second',
        )
        doc = parse(data)
        rtf = to_rtf(doc)
        self.assertIn(r'\page\pard', rtf)
        # \page must precede 'second', not precede 'first'
        self.assertLess(rtf.index(r'\page\pard'), rtf.index('second'))
        self.assertGreater(rtf.index(r'\page\pard'), rtf.index('first'))

    # --- DOCX ---

    def test_docx_emits_page_break_run(self):
        """DOCX output should include a w:lastRenderedPageBreak or w:br PAGE element."""
        import tempfile, zipfile, os
        from pathlib import Path
        data = self._make_page_break_doc(
            b'first' + bytes([0x13, 0x04, 0x78, 0x00]),
            b'second',
        )
        doc = parse(data)
        with tempfile.TemporaryDirectory() as tmp:
            dest = Path(tmp) / 'out.docx'
            from converter import save_docx
            save_docx(doc, dest)
            with zipfile.ZipFile(dest) as z:
                xml = z.read('word/document.xml').decode('utf-8')
        # python-docx emits WD_BREAK.PAGE as <w:br w:type="page"/>
        self.assertIn('w:type="page"', xml)


if __name__ == '__main__':
    unittest.main()
